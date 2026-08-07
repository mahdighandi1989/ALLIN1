"""AI document import: the .docx fast-path persists across the right customer
(no AI key needed), the JSON parser is tolerant, and multi-customer persist is
deduped (re-import never duplicates guarantors)."""
from io import BytesIO

from docx import Document

from app.models.customer import Customer
from app.services import doc_ingest


async def _poll(client, headers, job_id, tries: int = 50):
    """Poll an import job to completion. Jobs run inline under ``import_inline``,
    so this returns on the first read; the loop is just defensive."""
    import asyncio
    for _ in range(tries):
        r = await client.get(f"/api/imports/jobs/{job_id}", headers=headers)
        assert r.status_code == 200, r.text
        j = r.json()
        if j["status"] != "running":
            return j
        await asyncio.sleep(0.02)
    raise AssertionError(f"job {job_id} never finished")


def test_parse_model_json_tolerant():
    assert doc_ingest.parse_model_json('```json\n{"customers":[]}\n```') == {"customers": []}
    assert doc_ingest.parse_model_json('blah {"a":1} trailing')["a"] == 1
    assert doc_ingest.parse_model_json("not json") == {}


async def test_persist_customer_multi_and_dedup(db_session):
    payload = {
        "account_no": "2624-330011-1", "name": "Multi Co", "account_type": "corporate",
        "branch": "AL MAKTOUM - 2624",
        "fields": {"aecb_score": "640", "trade_license_no": "TL-9", "business_type": "Trading"},
        "guarantors": [{"name": "Ali Guarantor", "account": "330099", "branch": "2624"}],
        "review": {"date_of_review": "01/06/2026", "purpose": "fresh loan"},
    }
    r1 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r1["ok"] and r1["account_no"] == "330011"  # 6-digit core extracted
    assert "aecb_score" in r1["fields_saved"]
    assert r1["guarantors_added"] == 1
    await db_session.commit()

    # Re-persist the same → guarantor updated in place, not duplicated.
    r2 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r2["guarantors_added"] == 0 and r2["guarantors_updated"] == 1
    await db_session.commit()


def _draft_docx() -> bytes:
    d = Document()
    t = d.add_table(rows=0, cols=2)
    for k, v in [("Customer Name:", "EFCO TRADING LLC"), ("Account Number:", "2624 115524 011"),
                 ("Borrower Type", "SME / Corporate")]:
        c = t.add_row().cells
        c[0].text = k
        c[1].text = v
    d.add_paragraph("fresh commercial loan facility of AED 8,000,000/- for a period of 48 months. "
                    "Interest rate to be at 12% p.a. customer request letter dated.27/04/2026.")
    buf = BytesIO()
    d.save(buf)
    return buf.getvalue()


async def test_analyze_docx_fastpath(client, auth_headers, db_session, import_inline):
    db_session.add(Customer(account_no="115524", name="Old"))
    await db_session.commit()
    files = {"file": ("efco.docx", _draft_docx(),
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    r = await client.post("/api/imports/analyze", headers=auth_headers, files=files)
    assert r.status_code == 200, r.text
    job = await _poll(client, auth_headers, r.json()["job_id"])
    assert job["status"] == "done", job
    b = job["result"]
    assert b["ok"] and b["model"] == "Word draft parser"
    assert any(c.get("account_no") == "115524" for c in b["customers"])
    # data landed where other forms read it
    rd = await client.get("/api/crm/offer-letter-data/115524", headers=auth_headers)
    assert rd.status_code == 200
    rc = await client.get("/api/crm/credit-reviews/115524", headers=auth_headers)
    assert len(rc.json()) >= 1


def _xlsx_bytes() -> bytes:
    import io
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Account", "Name", "Nationality"])
    ws.append(["2624-440011-1", "Alpha Co", "Iran"])
    ws.append(["2624-440022-2", "Beta Co", "UAE"])
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()


def test_workbook_to_text_and_chunk():
    txt = doc_ingest.workbook_to_text(_xlsx_bytes(), "t.xlsx")
    assert "Alpha Co" in txt and "440022" in txt
    chunks = doc_ingest.chunk_text(txt, 50)
    assert len(chunks) >= 1
    # CSV path
    csv = doc_ingest.workbook_to_text(b"Account,Name\n440033,Gamma", "t.csv")
    assert "Gamma" in csv


async def test_analyze_xlsx_reaches_branch(client, auth_headers, import_inline):
    # No AI key in tests → the spreadsheet branch parses then reports no usable
    # model. The job captures that as an error with http_status 400/502, proving
    # the Excel path is wired (not a 415/500).
    files = {"file": ("book.xlsx", _xlsx_bytes(),
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
    r = await client.post("/api/imports/analyze", headers=auth_headers, files=files)
    assert r.status_code == 200, r.text
    job = await _poll(client, auth_headers, r.json()["job_id"])
    assert job["status"] == "error", job
    assert job["http_status"] in (400, 502), job


async def test_reupload_same_file_no_duplicate_review(client, auth_headers, db_session, import_inline):
    db_session.add(Customer(account_no="115524", name="X"))
    await db_session.commit()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    for _ in range(2):  # upload + extract the SAME file twice
        r = await client.post("/api/imports/analyze", headers=auth_headers,
                              files={"file": ("efco.docx", _draft_docx(), mime)})
        assert r.status_code == 200, r.text
        job = await _poll(client, auth_headers, r.json()["job_id"])
        assert job["status"] == "done", job
    rc = await client.get("/api/crm/credit-reviews/115524", headers=auth_headers)
    assert len(rc.json()) == 1  # deduped per review date — not duplicated on re-upload


async def test_persist_properties_dedup(db_session):
    base = {"account_no": "550011", "name": "Prop Co", "account_type": "corporate",
            "properties": [{"prop_type": "Land & Building", "address": "Deira", "city": "Dubai",
                            "plate_no": "P-9", "valuation": "15,400,000", "valuation_currency": "AED"}]}
    r1 = await doc_ingest.persist_customer(db_session, base, "tester")
    assert r1["properties_added"] == 1
    await db_session.commit()
    # same property (matched by plate_no) → updated, not duplicated
    r2 = await doc_ingest.persist_customer(db_session, base, "tester")
    assert r2["properties_added"] == 0 and r2["properties_updated"] == 1
    await db_session.commit()
    from app.models.profile_entities import MortgagedProperty
    from sqlalchemy import select as _sel
    rows = (await db_session.execute(_sel(MortgagedProperty).where(MortgagedProperty.account_no == "550011"))).scalars().all()
    assert len(rows) == 1 and float(rows[0].valuation) == 15400000


def test_extraction_is_schema_driven():
    """The field list the model is asked for is derived from the CustomerProfile
    schema — so previously-missed sub-fields are covered and a future column would
    be picked up automatically (no prompt edit)."""
    fields = doc_ingest.extractable_profile_fields()
    for k in ("visa_issue", "visa_type", "tenancy_issue", "tenancy_address",
              "emirates_id_golden", "nationality", "trade_license_issue", "auditor",
              "grade", "call_report", "previous_files"):
        assert k in fields, f"{k} should be extractable"
    # housekeeping / file-path / officer-notes columns are excluded
    for k in ("data_json", "passport_doc", "profile_completeness", "account_no",
              "passport_remarks", "created_at"):
        assert k not in fields, f"{k} should NOT be extractable"
    # and they actually reach the prompt sent to the model
    for k in ("visa_type", "tenancy_address", "emirates_id_golden"):
        assert k in doc_ingest.EXTRACTION_PROMPT


async def test_persist_full_kyc_schema_driven(db_session):
    """Every recognised KYC sub-field is promoted to its real column (incl. the ones
    the old hardcoded map dropped: visa type/issue, tenancy address, golden EID)."""
    payload = {
        "account_no": "660011", "name": "KYC Co", "account_type": "corporate",
        "fields": {
            "visa_no": "V-123", "visa_issue": "01/01/2024", "visa_expiry": "01/01/2027",
            "visa_type": "Investor", "tenancy_no": "T-9", "tenancy_issue": "02/02/2025",
            "tenancy_address": "Marina, Dubai", "emirates_id_golden": "Yes",
            "nationality": "Iran", "auditor": "KPMG", "monthly_salary": "45000",
        },
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"]
    await db_session.commit()
    from app.models.crm import CustomerProfile
    from sqlalchemy import select as _sel
    cp = (await db_session.execute(_sel(CustomerProfile).where(CustomerProfile.account_no == "660011"))).scalar_one()
    assert cp.visa_type == "Investor" and cp.visa_no == "V-123"
    assert cp.tenancy_address == "Marina, Dubai" and cp.tenancy_no == "T-9"
    assert cp.emirates_id_golden == "Yes"
    assert cp.passport_nationality == "Iran"   # via the "nationality" alias
    assert cp.auditor == "KPMG" and cp.monthly_salary == "45000"


async def test_persist_kyc_renewal_updates_date_but_fillempty_text(db_session):
    """KYC dates update on renewal (later wins); curated text is never clobbered."""
    from app.models.crm import CustomerProfile
    from sqlalchemy import select as _sel
    cp = CustomerProfile(account_no="660022", passport_expiry="01/01/2026",
                         passport_nationality="India")
    db_session.add(cp)
    await db_session.commit()
    payload = {"account_no": "660022", "name": "X", "fields": {
        "passport_expiry": "01/01/2031",   # renewal → later date wins
        "nationality": "Pakistan",          # already set → must NOT clobber
    }}
    await doc_ingest.persist_customer(db_session, payload, "tester")
    await db_session.commit()
    row = (await db_session.execute(_sel(CustomerProfile).where(CustomerProfile.account_no == "660022"))).scalar_one()
    assert row.passport_expiry == "01/01/2031"     # updated on renewal
    assert row.passport_nationality == "India"      # curated value preserved


async def test_persist_partners_and_facilities(db_session):
    """Import now also fills the corporate credit-file form: partners (Partner
    table) and facility records (Facility table) the form reads from."""
    from app.models.customer import Customer
    from app.models.profile_entities import Partner
    from app.models.facility import Facility
    from sqlalchemy import select as _sel
    payload = {
        "account_no": "770011", "name": "Corp Co", "account_type": "corporate",
        "partners": [{"name": "Ali Partner", "nationality": "Iran", "share": "60"},
                     {"name": "Sara Partner", "nationality": "UAE", "share": "40"}],
        "facilities": [{"facility_type": "overdraft", "amount": "500,000", "interest_rate": "8%",
                        "expiry_date": "31/12/2027"},
                       {"facility_type": "loan", "amount": "2000000"}],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["partners_added"] == 2 and r["facilities_added"] == 2
    await db_session.commit()
    cid = (await db_session.execute(_sel(Customer.id).where(Customer.account_no == "770011"))).scalar_one()
    parts = (await db_session.execute(_sel(Partner).where(Partner.account_no == "770011"))).scalars().all()
    assert {p.name for p in parts} == {"Ali Partner", "Sara Partner"}
    facs = (await db_session.execute(_sel(Facility).where(Facility.customer_id == cid))).scalars().all()
    od = next(f for f in facs if str(getattr(f.facility_type, "value", f.facility_type)) == "overdraft")
    assert float(od.amount) == 500000 and float(od.interest_rate) == 8
    assert od.expiry_date is not None

    # Re-import the same → no duplicate partners/facilities (deduped, fill-empty).
    r2 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r2["partners_added"] == 0 and r2["facilities_added"] == 0
    await db_session.commit()


async def test_persist_security_and_summary_fields(db_session):
    """The credit-file summary header (grade/call report/previous files/undertaking)
    and the Security matrix are extracted and persisted (columns + data_json)."""
    import json as _j
    from app.models.crm import CustomerProfile
    from sqlalchemy import select as _sel
    payload = {
        "account_no": "880011", "name": "Sec Co", "account_type": "corporate",
        "fields": {"grade": "GOOD", "call_report": "done 01/2026", "previous_files": "2",
                   "undertaking_from": "Guarantor/s"},
        "security": [{"type": "Cheques", "for_facility": "overdraft", "aed": "100000"},
                     {"type": "Collaterals", "aed": "5000000"}],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["security_added"] == 2
    await db_session.commit()
    cp = (await db_session.execute(_sel(CustomerProfile).where(CustomerProfile.account_no == "880011"))).scalar_one()
    assert cp.grade == "GOOD" and cp.call_report == "done 01/2026" and cp.previous_files == "2"
    assert cp.undertaking_from == "Guarantor/s"
    sd = _j.loads(cp.data_json)["security_details"]
    assert {row["type"] for row in sd} == {"Cheques", "Collaterals"}
    chq = next(row for row in sd if row["type"] == "Cheques")
    assert chq["aed"] == "100000" and chq["for_facility"] == "overdraft"
    # re-import the same security → no duplicate rows
    r2 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r2["security_added"] == 0


async def test_persist_facility_rate_overflow_is_dropped(db_session):
    """A mis-extracted interest_rate (3190 into a Numeric(5,2) column) must NOT
    crash the import — it is dropped, the facility still saves with its amount."""
    from app.models.customer import Customer
    from app.models.facility import Facility
    from sqlalchemy import select as _sel
    payload = {
        "account_no": "990011", "name": "Big Co", "account_type": "corporate",
        "facilities": [{"facility_type": "overdraft", "amount": "1000000",
                        "interest_rate": "3190", "expiry_date": "30/05/2027"}],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["facilities_added"] == 1
    await db_session.commit()
    cid = (await db_session.execute(_sel(Customer.id).where(Customer.account_no == "990011"))).scalar_one()
    fac = (await db_session.execute(_sel(Facility).where(Facility.customer_id == cid))).scalar_one()
    assert float(fac.amount) == 1000000
    assert fac.interest_rate is None  # the out-of-range rate was dropped, not stored


async def test_partner_reconcile_and_impossible_pct(db_session):
    """Same partner spelled differently collapses to ONE row; an impossible share
    (a capital amount as %) is dropped, not stored."""
    from app.models.profile_entities import Partner
    from sqlalchemy import select as _sel
    await doc_ingest.persist_customer(db_session, {
        "account_no": "771100", "name": "Co",
        "partners": [{"name": "Yousef Mohamed Alhammadi", "share": "11.000000%"},
                     {"name": "Capital Holder", "share": "3300000%"}],
    }, "t")
    await db_session.commit()
    # second document: same person, extra middle name → must NOT duplicate
    await doc_ingest.persist_customer(db_session, {
        "account_no": "771100", "name": "Co",
        "partners": [{"name": "Yousef Mohamed Ahmed Alhammadi", "nationality": "UAE"}],
    }, "t")
    await db_session.commit()
    rows = (await db_session.execute(_sel(Partner).where(
        Partner.account_no == "771100", Partner.is_deleted == False))).scalars().all()
    yousefs = [r for r in rows if "ousef" in r.name]
    assert len(yousefs) == 1, [r.name for r in rows]   # collapsed across spellings
    assert yousefs[0].share == "11%"                    # trailing zeros trimmed
    cap = next(r for r in rows if r.name == "Capital Holder")
    assert (cap.share or "") == ""                       # impossible % dropped


def test_same_person_and_clean_pct_units():
    assert doc_ingest._same_person("Yousef Alhammadi", "Yousef Mohamed Alhammadi")
    assert not doc_ingest._same_person("Khaled Shojae", "Valid Shojae")
    assert doc_ingest._clean_pct("45.0000000%") == "45%"
    assert doc_ingest._clean_pct("3300000%") == ""
    assert doc_ingest._clean_pct("33.33") == "33.33%"


async def test_property_consolidates_same_number_one_import(db_session):
    """The same property described several ways in ONE file (number stuffed in the
    address, different type labels) collapses to a single row with merged details."""
    from app.models.profile_entities import MortgagedProperty
    from sqlalchemy import select as _sel
    await doc_ingest.persist_customer(db_session, {
        "account_no": "453100", "name": "Co",
        "properties": [
            {"prop_type": "Property", "address": "Property no. 638/140", "city": "Shiraz"},
            {"prop_type": "FLAT", "address": "Property no. 638/140", "valuation": "458000", "valuation_currency": "AED"},
            {"prop_type": "Apartment", "address": "Property no. 638/140"},
        ],
    }, "t")
    await db_session.commit()
    rows = (await db_session.execute(_sel(MortgagedProperty).where(
        MortgagedProperty.account_no == "453100", MortgagedProperty.is_deleted == False))).scalars().all()
    assert len(rows) == 1, [r.prop_type for r in rows]   # consolidated to ONE
    assert float(rows[0].valuation) == 458000             # detail merged in


def test_prop_token_extracts_number():
    assert doc_ingest._prop_token("Property no. 638/140") == "638/140"
    assert doc_ingest._prop_token("FLAT", "", "deed 638 / 140") == "638/140"
    assert doc_ingest._prop_token("no number here") == ""


async def test_fail_orphaned_jobs_marks_error(db_session, monkeypatch):
    """A job left 'running' by a killed process is marked errored at startup so the
    browser's poll stops waiting; finished jobs are untouched."""
    from contextlib import asynccontextmanager
    from app.routers import imports as imp
    from app.models.import_job import ImportJob
    from sqlalchemy import select as _sel
    db_session.add(ImportJob(id="orphan1", status="running", filename="x.pdf"))
    db_session.add(ImportJob(id="done1", status="done", filename="y.pdf"))
    await db_session.commit()

    @asynccontextmanager
    async def _reuse():
        yield db_session
    monkeypatch.setattr(imp, "_job_session", _reuse)

    n = await imp.fail_orphaned_jobs()
    assert n == 1
    orphan = (await db_session.execute(_sel(ImportJob).where(ImportJob.id == "orphan1"))).scalar_one()
    assert orphan.status == "error" and orphan.http_status == 503
    done = (await db_session.execute(_sel(ImportJob).where(ImportJob.id == "done1"))).scalar_one()
    assert done.status == "done"  # untouched


async def test_kyc_date_field_rejects_non_date(db_session):
    """A non-date in an issue/expiry field (e.g. the place of issue 'ABU DHABI')
    is dropped, not stored; a real date is normalised and kept."""
    from app.models.crm import CustomerProfile
    from sqlalchemy import select as _sel
    await doc_ingest.persist_customer(db_session, {
        "account_no": "350994", "name": "Azoora",
        "fields": {"trade_license_issue": "ABU DHABI", "trade_license_expiry": "10/03/2027"},
    }, "t")
    await db_session.commit()
    cp = (await db_session.execute(_sel(CustomerProfile).where(
        CustomerProfile.account_no == "350994"))).scalar_one()
    assert (cp.trade_license_issue or "") == ""        # place-of-issue dropped
    assert cp.trade_license_expiry == "10/03/2027"      # real date kept


def test_merge_customer_reassembles_lists_across_chunks():
    """A record split across PDF chunks (type in one, details in another) is
    field-merged, not dropped or duplicated."""
    into = {"properties": [{"prop_type": "Land & Building"}], "partners": [{"name": "Ali"}]}
    more = {"properties": [{"prop_type": "Land & Building", "address": "Deira", "valuation": "15000000"}],
            "partners": [{"name": "Sara", "nationality": "Iran"}],
            "facilities": [{"facility_type": "overdraft", "amount": "500000"}]}
    doc_ingest.merge_customer(into, more)
    assert len(into["properties"]) == 1  # field-merged, not duplicated
    assert into["properties"][0]["address"] == "Deira" and into["properties"][0]["valuation"] == "15000000"
    assert len(into["partners"]) == 2
    assert into["facilities"][0]["facility_type"] == "overdraft"


def test_split_pdf_and_offset():
    import io
    import pytest
    PdfWriter = pytest.importorskip("pypdf").PdfWriter  # prod-only dep
    w = PdfWriter()
    for _ in range(5):
        w.add_blank_page(width=200, height=200)
    buf = io.BytesIO(); w.write(buf)
    chunks, n = doc_ingest.split_pdf(buf.getvalue(), max_bytes=10_000_000, max_pages=2)
    assert n == 5 and len(chunks) == 3 and chunks[0][0] == 1 and chunks[1][0] == 3
    # the streaming generator yields the same chunks one at a time (low memory)
    gen = list(doc_ingest.pdf_chunks(buf.getvalue(), max_bytes=10_000_000, max_pages=2))
    assert [c[0] for c in gen] == [1, 3, 5]
    assert doc_ingest.offset_pages("1-2", 2) == "3-4"


async def test_persist_facilities_deposit_guard_and_required_securities(db_session, client, auth_headers):
    """The owner's real regression: the sanction's FD-underlien text came back
    as a third 'facility' (type → other) and surfaced on the Offer Letter as a
    phantom «Credit Facility» row; and the REQUIRED SECURITIES list never
    reached the letter. The guard drops deposit/summary pseudo-facilities,
    required_securities lands on the profile, re-import stays deduped, and
    /offer-letter-data serves both (proper type labels included)."""
    payload = {
        "account_no": "2900-301408-010", "name": "Al Ain Glass & Mirrors", "account_type": "corporate",
        "facilities": [
            {"facility_type": "overdraft", "amount": "3500000", "interest_rate": "5.25"},
            {"facility_type": "cheque_discounting", "amount": "2800000", "interest_rate": "11"},
            # the junk that caused the phantom row — a deposit posing as a facility
            {"facility_type": "credit facility", "amount": "3500000",
             "notes": "Fixed Deposit 365 days, Ref AJMN FD-2025-73, start 29NOV25"},
        ],
        "required_securities": "Signed undertaking form for total facility amount from the borrower.\nRenewed letter of lien and authority to set off for advances against fixed deposit amounting to AED 3,500,000/- held underlien in same account.",
    }
    r1 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r1["ok"]
    assert r1["facilities_added"] == 2                    # OD + CD only
    assert r1["facilities_skipped_deposits"] == 1         # the FD junk was dropped
    await db_session.commit()

    # re-import the SAME sanction → matched per type, nothing duplicated
    r2 = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r2["facilities_added"] == 0 and r2["facilities_updated"] == 2
    assert r2["facilities_skipped_deposits"] == 1
    await db_session.commit()

    # the Offer Letter sees exactly TWO rows, proper labels, and the sanction's securities text
    r = await client.get("/api/crm/offer-letter-data/301408", headers=auth_headers)
    assert r.status_code == 200, r.text
    body = r.json()
    facs = body["Facilities"]
    assert len(facs) == 2
    assert {f["type"] for f in facs} == {"Overdraft", "Cheque Discount"}   # no raw snake_case
    assert "held underlien" in body["RequiredSecurities"]


async def test_analyze_xlsx_timeout_retries_then_succeeds(client, auth_headers, db_session, import_inline, monkeypatch):
    """A slow model minute must not fail the import: the Excel path uses the
    long extraction deadline (180s) and retries a timed-out call ONCE."""
    from app.ai import inference as inf

    calls = {"n": 0, "timeouts": []}

    async def fake_complete(db, prompt, **kw):
        calls["n"] += 1
        calls["timeouts"].append(kw.get("timeout"))
        if calls["n"] == 1:
            return {"ok": False, "error": "timed out after 180s", "text": "", "model": "M"}
        return {"ok": True, "model": "M", "error": None,
                "text": '{"customers": [{"account_no": "440022", "name": "Alpha Co", "fields": {}}], "documents": []}'}

    monkeypatch.setattr(inf, "complete", fake_complete)
    files = {"file": ("book.xlsx", _xlsx_bytes(),
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
    r = await client.post("/api/imports/analyze", headers=auth_headers, files=files)
    assert r.status_code == 200, r.text
    job = await _poll(client, auth_headers, r.json()["job_id"])
    assert job["status"] == "done", job
    assert calls["n"] == 2                              # timed out once, retried, succeeded
    assert all(t == 180.0 for t in calls["timeouts"])   # long extraction deadline on both


def test_table_fallback_customers_maps_accounts_and_properties():
    """Deterministic safety net for account-column spreadsheets (v49): rows
    group per 6-digit account; property columns map by their Persian headers."""
    txt = ("## Sheet: 2624\n"
           "شماره,شعبه,شماره حساب,نام مشتری,شماره سند/ پلاک ثبتی,شهر,نوع,وضعیت ملک\n"
           "1,2624,111150,ALL CHEM  INTL,25/28,شیراز,باغ,در رهن می باشد\n"
           "2,2624,111150,ALL CHEM  INTL,25/29,شیراز,ویلا,فک رهن شده است\n"
           "3,2624,111714,PETROPOL,5128,تهران,ویلا,در رهن می باشد\n")
    cs = doc_ingest.table_fallback_customers(txt)
    assert {c["account_no"] for c in cs} == {"111150", "111714"}
    c1 = next(c for c in cs if c["account_no"] == "111150")
    assert c1["name"] == "ALL CHEM INTL"
    assert len(c1["properties"]) == 2
    assert c1["properties"][0]["mortgage_deed_no"] == "25/28"
    assert c1["properties"][1]["remarks"] == "فک رهن شده است"
    # a sheet with no account column yields nothing (no guessing)
    assert doc_ingest.table_fallback_customers("## Sheet: x\na,b\n1,2\n") == []


async def test_analyze_xlsx_model_garbage_falls_back_to_table(client, auth_headers, db_session, import_inline, monkeypatch):
    """When the model answers but returns nothing parseable, the deterministic
    table parser still imports the accounts instead of failing the job."""
    from app.ai import inference as inf

    async def fake_complete(db, prompt, **kw):
        return {"ok": True, "model": "M", "error": None, "text": "sorry, I cannot help with that"}

    monkeypatch.setattr(inf, "complete", fake_complete)
    files = {"file": ("book.xlsx", _xlsx_bytes(),
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
    r = await client.post("/api/imports/analyze", headers=auth_headers, files=files)
    assert r.status_code == 200, r.text
    job = await _poll(client, auth_headers, r.json()["job_id"])
    assert job["status"] == "done", job
    assert "قطعی" in str(job.get("result", {}).get("model") or job.get("result", ""))  # fallback engaged


async def test_letter_attachment_excel_falls_back_to_table(db_session, monkeypatch):
    """The letter-assistant attachment extraction gets the same deterministic
    table safety net as the Import page (v50)."""
    from app.ai import inference as inf
    from app.services import letter_attachment_extract as lax

    async def fake_complete(db, prompt, **kw):
        return {"ok": True, "model": "M", "error": None, "text": "cannot extract"}

    monkeypatch.setattr(inf, "complete", fake_complete)
    res = await lax.extract_attachment(
        db_session, data=_xlsx_bytes(), filename="book.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        letter_ctx={})
    assert res["ok"], res
    accs = {c["account_no"] for c in res["customers"]}
    assert "440022" in accs
    assert "قطعی" in str(res.get("model"))


# ---------------- v58: person-identity + full property fields ----------------

def test_extraction_prompt_asks_for_identity_and_property_details():
    """Owner requirement (v58): manager/partner national IDs + passports with
    dates, guarantor national ID, property owner + owner national ID, postal
    code, areas, and the insurance policy number + computer code + issue/expiry
    must all be asked of the model."""
    p = doc_ingest.EXTRACTION_PROMPT
    for k in ("\"role\"", "\"national_id\"", "\"passport_issue\"", "\"passport_expiry\"",
              "\"emirates_id_no\"", "\"emirates_id_expiry\"",
              "\"owner\"", "\"owner_national_id\"", "\"postal_code\"",
              "\"land_area\"", "\"infra_area\"", "\"building_age\"", "\"zone\"",
              "\"insurance_no\"", "\"insurance_computer_code\"", "\"insurance_issue\"",
              "\"last_valuation_date\""):
        assert k in p, k
    # the guidance rules for these facts reach the model too
    assert "PERSON IDENTITY NUMBERS MATTER" in p
    assert "کد ملی" in p
    # the customer's own national_id column is schema-driven into the prompt
    assert "national_id" in doc_ingest.extractable_profile_fields()


async def test_persist_person_identity_and_property_details(db_session):
    payload = {
        "account_no": "701001", "name": "Identity Co LLC", "account_type": "corporate",
        "fields": {"national_id": "0012345678"},
        "partners": [{"name": "Ali Rezaei", "role": "Manager", "nationality": "Iranian",
                      "national_id": "0087654321", "passport_no": "P1234567",
                      "passport_issue": "01/01/2020", "passport_expiry": "01/01/2030",
                      "emirates_id_no": "784-1980-1234567-1", "emirates_id_expiry": "05/05/2027",
                      "share": "50"}],
        "guarantors": [{"name": "Hasan Karimi", "account": "701002", "national_id": "0055512345"}],
        "properties": [{"prop_type": "Villa", "address": "Street 5, Dubai", "city": "Dubai",
                        "postal_code": "12345", "owner": "Ali Rezaei", "owner_national_id": "0087654321",
                        "land_area": "500", "infra_area": "320", "building_age": "8", "zone": "A",
                        "valuation": "2500000", "last_valuation_date": "10/06/2026",
                        "mortgage_amount": "1500000", "plate_no": "638/140", "mortgage_deed_no": "12345",
                        "insurance_no": "INS-991", "insurance_computer_code": "CC-77821",
                        "insurance_issue": "01/03/2026", "insurance_expiry": "01/03/2027"}],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"]
    await db_session.commit()

    from app.models.crm import CustomerProfile
    from app.models.profile_entities import Partner, MortgagedProperty
    from app.models.guarantor import Guarantor
    from sqlalchemy import select

    cp = (await db_session.execute(select(CustomerProfile).where(
        CustomerProfile.account_no == "701001"))).scalar_one()
    assert cp.national_id == "0012345678"

    pt = (await db_session.execute(select(Partner).where(
        Partner.account_no == "701001"))).scalars().one()
    assert pt.role == "Manager" and pt.national_id == "0087654321"
    assert pt.passport_no == "P1234567" and pt.passport_issue == "01/01/2020"
    assert pt.passport_expiry == "01/01/2030"
    assert pt.emirates_id_no == "784-1980-1234567-1" and pt.emirates_id_expiry == "05/05/2027"

    g = (await db_session.execute(select(Guarantor).where(
        Guarantor.account_no == "701001"))).scalars().one()
    assert g.national_id == "0055512345"

    pr = (await db_session.execute(select(MortgagedProperty).where(
        MortgagedProperty.account_no == "701001"))).scalars().one()
    assert pr.owner == "Ali Rezaei" and pr.owner_national_id == "0087654321"
    assert pr.postal_code == "12345" and pr.land_area == "500" and pr.infra_area == "320"
    assert pr.building_age == "8" and pr.zone == "A"
    assert pr.insurance_no == "INS-991" and pr.insurance_computer_code == "CC-77821"
    assert pr.insurance_issue == "01/03/2026" and pr.insurance_expiry == "01/03/2027"
    assert pr.last_valuation_date == "10/06/2026"
    assert float(pr.valuation) == 2500000.0 and float(pr.mortgage_amount) == 1500000.0


async def test_attachment_text_deterministic_paths(db_session):
    """v61: the full_check pass reads attachments as TEXT without any model or
    DB write — Excel/CSV and plain text are fully deterministic."""
    from app.services import letter_attachment_extract as lax
    r = await lax.attachment_text(db_session, data=_xlsx_bytes(), filename="t.xlsx",
                                  mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert r["ok"] and "440022" in r["text"]
    r2 = await lax.attachment_text(db_session, data="سلام دنیا".encode("utf-8"),
                                   filename="a.txt", mimetype="text/plain")
    assert r2["ok"] and "سلام" in r2["text"]
    r3 = await lax.attachment_text(db_session, data=b"junk", filename="x.zzz", mimetype="")
    assert not r3["ok"]


# ---------------- v64: property EVENT TIMELINE (several valuations, …) ----------------

def test_extraction_prompt_asks_for_property_event_history():
    p = doc_ingest.EXTRACTION_PROMPT
    assert "PROPERTY EVENT HISTORY" in p
    for k in ("\"events\"", "remortgage", "additional_mortgage", "release", "فک رهن", "ترهین مجدد"):
        assert k in p, k
    assert "report them ALL" in p            # several valuations, never only the latest


async def test_persist_property_event_timeline_and_dedup(db_session):
    payload = {
        "account_no": "702001", "name": "Timeline Co", "account_type": "corporate",
        "properties": [{
            "prop_type": "Villa", "address": "تهران، خیابان آزادی", "city": "تهران", "country": "Iran",
            "plate_no": "638/140", "mortgage_deed_no": "12345",
            "valuation": "9000000", "last_valuation_date": "1403/05/12",
            "events": [
                {"event_type": "valuation", "date": "1399/02/10", "amount": "5000000", "currency": "IRR"},
                {"event_type": "valuation", "date": "1401/07/03", "amount": "7500000", "currency": "IRR"},
                {"event_type": "valuation", "date": "1403/05/12", "amount": "9000000", "currency": "IRR"},
                {"event_type": "mortgage", "date": "1399/03/01", "amount": "4000000", "currency": "IRR"},
                {"event_type": "ترهین مجدد", "date": "1401/08/15", "amount": "6000000"},  # normalized → other? no: not in set → other
                {"event_type": "remortgage", "date": "1401/08/15", "amount": "6000000", "currency": "IRR"},
                {"event_type": "release", "date": "1404/01/20"},
                {"event_type": "junk"},   # neither date nor amount → dropped
            ],
        }],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["property_events_added"] >= 6
    await db_session.commit()

    from app.models.profile_entities import PropertyEvent
    from sqlalchemy import select
    evs = (await db_session.execute(select(PropertyEvent).where(
        PropertyEvent.account_no == "702001"))).scalars().all()
    types = sorted(e.event_type for e in evs)
    # all three valuations kept as separate dated rows
    assert types.count("valuation") == 3
    assert "mortgage" in types and "remortgage" in types and "release" in types
    rel = next(e for e in evs if e.event_type == "release")
    assert rel.event_date == "1404/01/20"

    # re-import the SAME payload → nothing doubles
    r2 = await doc_ingest.persist_customer(db_session, payload, "tester")
    await db_session.commit()
    assert r2["property_events_added"] == 0
    evs2 = (await db_session.execute(select(PropertyEvent).where(
        PropertyEvent.account_no == "702001"))).scalars().all()
    assert len(evs2) == len(evs)


def test_build_facts_includes_property_history_generically():
    """The letter/AI facts now carry the properties (ALL columns, schema-driven)
    plus each property's event timeline — so any current or FUTURE property
    column is available wherever the DB facts are fetched."""
    from app.services import letter_assistant as la
    from app.models.profile_entities import MortgagedProperty, PropertyEvent

    prop = MortgagedProperty(id="P1", account_no="702002", plate_no="638/140",
                             city="Dubai", country="UAE", owner="Ali",
                             owner_national_id="0012345678", postal_code="12345",
                             valuation=2500000, last_valuation_date="10/06/2026")
    ev = PropertyEvent(id="E1", property_id="P1", account_no="702002",
                       event_type="valuation", event_date="10/06/2026",
                       amount=2500000, currency="AED")
    ev2 = PropertyEvent(id="E2", property_id="P1", account_no="702002",
                        event_type="release", event_date="01/01/2027")
    facts = la.build_facts(None, {}, [], [], properties=[prop],
                           property_events=[ev, ev2])
    p = facts["properties"][0]
    assert p["plate_no"] == "638/140" and p["owner_national_id"] == "0012345678"
    assert p["postal_code"] == "12345"          # v58 column flows automatically
    hist = p["history"]
    assert {h["event_type"] for h in hist} == {"valuation", "release"}
    assert any(h.get("date") == "01/01/2027" for h in hist)


async def test_import_persists_irr_cheque_details(db_session):
    """v99 — guarantor entries extracted from sanction documents carry the IRR
    cheque details through to the Guarantor row (fill-empty)."""
    payload = {
        "account_no": "2624-115529-011", "name": "SATIN STAR TRADING LLC",
        "account_type": "corporate",
        "guarantors": [{
            "name": "M/s Hani pokht iraninan", "cheque_no": "536001/237986",
            "cheque_currency": "IRR", "irr_amount": "6383360000000",
            "irr_rate": "398960", "coverage_pct": "200",
            "issuer_bank": "karafarin bank - 5300114",
        }],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["guarantors_added"] == 1
    await db_session.commit()
    from app.models.guarantor import Guarantor
    from sqlalchemy import select
    row = (await db_session.execute(select(Guarantor).where(Guarantor.account_no == "115529"))).scalars().first()
    assert row.cheque_no == "536001/237986" and row.cheque_currency == "IRR"
    assert row.irr_amount == "6383360000000" and row.irr_rate == "398960"
    assert row.coverage_pct == "200" and "karafarin" in row.issuer_bank_code
    # the extraction prompt teaches the model to capture these fields
    assert "irr_amount" in doc_ingest.EXTRACTION_PROMPT and "IRR" in doc_ingest.EXTRACTION_PROMPT


async def test_import_links_cheque_to_facility_and_tenor(db_session):
    """v101 — the extracted facility_ref lands on Guarantor.facility_id, and a
    facility's tenor_months persists (both fill-empty), so the voucher form can
    suggest/auto-fill the IRR narrative from the DB."""
    payload = {
        "account_no": "115530", "name": "TENOR TEST LLC",
        "account_type": "corporate",
        "guarantors": [{
            "name": "M/s Issuer Co", "cheque_no": "990011",
            "cheque_currency": "IRR", "irr_amount": "1000000000",
            "facility_ref": "STF1260603000001",
        }],
        "facilities": [{"facility_type": "loan", "amount": "8000000",
                        "currency": "AED", "tenor_months": "48"}],
    }
    r = await doc_ingest.persist_customer(db_session, payload, "tester")
    assert r["ok"] and r["guarantors_added"] == 1 and r["facilities_added"] == 1
    await db_session.commit()
    from app.models.guarantor import Guarantor
    from app.models.facility import Facility
    from sqlalchemy import select
    g = (await db_session.execute(select(Guarantor).where(Guarantor.account_no == "115530"))).scalars().first()
    assert g.facility_id == "STF1260603000001"
    f = (await db_session.execute(select(Facility))).scalars().all()
    mine = [x for x in f if x.tenor_months == "48" and float(x.amount) == 8000000.0]
    assert mine, "tenor_months must persist on the imported facility"
    # the prompt asks for both links
    assert "facility_ref" in doc_ingest.EXTRACTION_PROMPT and "tenor_months" in doc_ingest.EXTRACTION_PROMPT
