"""Dropping a filled committee-approval draft (.docx) into the Offer Letter must
extract its fields, persist them to the customer record, and — crucially — never
create duplicates when the same draft is extracted again (or the Offer Letter is
later printed).
"""
from io import BytesIO

from docx import Document

from app.models.customer import Customer


def _draft_bytes() -> bytes:
    """A minimal draft that mimics the real مصوبه labels + prose."""
    doc = Document()
    t = doc.add_table(rows=0, cols=2)
    for label, val in [
        ("Customer Name:", "NAEIMEH GHARABI HASHEMI"),
        ("Account Number:", "2624 127987 006"),
        ("Borrower Type", "Retail"),
        ("Business Activity:", "Salaried"),
        ("Proposed Customer Rating", "B*"),
    ]:
        r = t.add_row().cells
        r[0].text = label
        r[1].text = val
    doc.add_paragraph(
        "Fresh Personal loan of AED 80,000/- is requested for a period of 48 months. "
        "Interest rate to be at 12% p.a. Based on customer letter dtd.11/06/2026."
    )
    gt = doc.add_table(rows=1, cols=3)
    c = gt.rows[0].cells
    c[0].text = "Muhammad Ebrahim (Guarantor -1)"
    c[1].text = "2624"
    c[2].text = "124076"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _files():
    return {"file": ("draft.docx", _draft_bytes(),
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}


async def test_extract_draft_fills_and_dedupes(client, auth_headers, db_session):
    db_session.add(Customer(account_no="127987", name="Old Name"))
    await db_session.commit()

    r1 = await client.post("/api/crm/extract-draft", headers=auth_headers,
                           files=_files(), data={"account_no": "127987"})
    assert r1.status_code == 200, r1.text
    b1 = r1.json()
    assert b1["account_type"] == "retail"
    assert b1["offer"]["LoanAmount"] == "80,000"
    assert b1["offer"]["FacilityType"] == "Personal Loan"
    assert b1["offer"]["LoanTenor"] == "48"
    assert b1["guarantors_added"] == 1

    # Re-extract the SAME draft → guarantor is updated in place, not duplicated.
    r2 = await client.post("/api/crm/extract-draft", headers=auth_headers,
                           files=_files(), data={"account_no": "127987"})
    assert r2.status_code == 200
    assert r2.json()["guarantors_added"] == 0
    assert r2.json()["guarantors_updated"] == 1

    rg = await client.get("/api/crm/guarantors/127987", headers=auth_headers)
    assert len(rg.json()) == 1

    # Extracted facts persisted to the shared profile (readable by other forms).
    rd = await client.get("/api/crm/offer-letter-data/127987", headers=auth_headers)
    assert rd.status_code == 200
    assert rd.json()["AccountType"] == "retail"
    assert rd.json()["Saved"].get("aecb_score") is None  # offer snapshot is separate
