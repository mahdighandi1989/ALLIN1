"""The AI attachment GENERATOR: the model only proposes a strict JSON spec;
these tests lock the deterministic gate (spec parsing/clamps), the server-side
renderers (real xlsx/docx round-trips), the endpoint wiring with a mocked
model, and the AI_GENERATED marker that default-excludes these files from the
extraction tool (their data came out of the database in the first place)."""
import io
import json

import pytest

from app.models.ai_config import AIModel, AIProvider
from app.services import letter_attachment_generate as gen


# ---------------- spec parsing / clamps ----------------

def test_parse_spec_excel_clamps_and_sanitizes():
    raw = json.dumps({
        "kind": "excel", "filename": 'گزارش/تسهیلات: <ویژه>؟',
        "title": "ت" * 500,
        "warnings": ["نرخ سود در پایگاه‌داده نبود", ""],
        "sheets": [{
            "name": "برگه[1]*", "columns": [f"ستون{i}" for i in range(50)],
            "rows": [[i, None, "x" * 900] for i in range(1000)],
        }],
    }, ensure_ascii=False)
    spec, warnings = gen.parse_spec(raw)
    assert spec["kind"] == "excel"
    assert "/" not in spec["filename"] and "<" not in spec["filename"]
    assert len(spec["title"]) == 200
    assert warnings == ["نرخ سود در پایگاه‌داده نبود"]
    sh = spec["sheets"][0]
    assert len(sh["columns"]) == gen.MAX_COLS
    assert len(sh["rows"]) == gen.MAX_ROWS
    assert sh["rows"][0][1] == ""                      # None → empty, never invented
    assert len(sh["rows"][0][2]) == gen.MAX_CELL
    assert "[" not in sh["name"] and "*" not in sh["name"]


def test_parse_spec_word_and_kind_inference():
    raw = json.dumps({
        "filename": "توضیحات",
        "paragraphs": ["بند سادهٔ متنی", {"text": "سرفصل", "heading": True},
                       {"text": "بند تراز", "align": "bogus"}],
    }, ensure_ascii=False)
    spec, _ = gen.parse_spec(raw)
    assert spec["kind"] == "word"                       # inferred from paragraphs
    assert spec["paragraphs"][0]["text"] == "بند سادهٔ متنی"
    assert spec["paragraphs"][1]["heading"] is True
    assert spec["paragraphs"][2]["align"] == "justify"  # bogus align → safe default


def test_parse_spec_rejects_garbage():
    with pytest.raises(ValueError):
        gen.parse_spec("no json here")
    with pytest.raises(ValueError):
        gen.parse_spec(json.dumps({"kind": "excel", "sheets": []}))
    with pytest.raises(ValueError):
        gen.parse_spec(json.dumps({"kind": "word", "paragraphs": [{"text": ""}]}))


# ---------------- renderers (real file round-trips) ----------------

def test_render_excel_roundtrip_rtl_and_styling():
    spec, _ = gen.parse_spec(json.dumps({
        "kind": "excel", "filename": "جدول آزمون", "title": "وضعیت املاک",
        "sheets": [{"name": "املاک", "columns": ["ردیف", "ملک", "ملاحظات"],
                    "rows": [["۱", "پلاک ۱۲", ""], ["۲", "پلاک ۱۵", "بیمه‌نامه ندارد"]]}],
    }, ensure_ascii=False))
    data, filename, mimetype = gen.render(spec)
    assert filename.endswith(".xlsx") and "spreadsheet" in mimetype
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data))
    ws = wb["املاک"]
    assert ws.sheet_view.rightToLeft is True
    assert ws.cell(row=1, column=1).value == "وضعیت املاک"     # merged title row
    assert ws.cell(row=2, column=1).value == "ردیف"            # styled header
    assert ws.cell(row=2, column=1).font.bold is True
    assert ws.cell(row=3, column=2).value == "پلاک ۱۲"
    assert ws.cell(row=4, column=3).value == "بیمه‌نامه ندارد"


def test_render_word_roundtrip_rtl():
    spec, _ = gen.parse_spec(json.dumps({
        "kind": "word", "filename": "توضیحات", "title": "توضیحات تکمیلی",
        "paragraphs": [{"text": "سرفصل اول", "heading": True},
                       {"text": "متن رسمی مطابق لحن نامه.", "align": "justify"}],
    }, ensure_ascii=False))
    data, filename, mimetype = gen.render(spec)
    assert filename.endswith(".docx") and "wordprocessing" in mimetype
    import docx
    d = docx.Document(io.BytesIO(data))
    texts = [p.text for p in d.paragraphs]
    assert "توضیحات تکمیلی" in texts and "سرفصل اول" in texts
    # every paragraph is bidi (RTL) — the whole point for a Persian letter
    assert all(p._p.xpath("./w:pPr/w:bidi") for p in d.paragraphs if p.text)


# ---------------- need_data protocol (cross-customer datasets) ----------------

def test_parse_need_data():
    ok = gen.parse_need_data(json.dumps({
        "need_data": {"datasets": ["properties", "bogus", "customers"], "branch": "Sheikh Zayed"},
    }))
    assert ok == {"datasets": ["properties", "customers"], "branch": "Sheikh Zayed"}
    # a spec reply is NOT a data request
    assert gen.parse_need_data(json.dumps({"kind": "excel", "sheets": []})) is None
    # only-bogus datasets → not a valid request
    assert gen.parse_need_data(json.dumps({"need_data": {"datasets": ["nope"]}})) is None
    assert gen.parse_need_data("garbage") is None


async def _seed_branch_data(db_session):
    from app.models.customer import Customer
    from app.models.profile_entities import MortgagedProperty

    db_session.add(Customer(account_no="ACC1", name="شرکت الف", branch="Sheikh Zayed",
                            relationship_manager="آقای مدیری"))
    db_session.add(Customer(account_no="ACC2", name="شرکت ب", branch="Deira",
                            relationship_manager="خانم دیگری"))
    db_session.add(MortgagedProperty(id="P1", account_no="ACC1", plate_no="1234/56",
                                     mortgage_deed_no="MD-77", insurance_no="INS-9",
                                     insurance_expiry="2026-09-01", city="دبی"))
    db_session.add(MortgagedProperty(id="P2", account_no="ACC2", plate_no="9999/99",
                                     mortgage_deed_no="MD-88"))
    await db_session.commit()


async def test_fetch_datasets_branch_filter(db_session):
    await _seed_branch_data(db_session)
    data, warnings = await gen.fetch_datasets(db_session, ["properties", "customers"], "sheikh zayed")
    assert warnings == []
    props = data["properties"]
    assert len(props) == 1 and props[0]["plate_no"] == "1234/56"
    assert props[0]["mortgage_deed_no"] == "MD-77"
    assert props[0]["insurance_no"] == "INS-9" and props[0]["insurance_expiry"] == "2026-09-01"
    assert props[0]["account_manager"] == "آقای مدیری" and props[0]["branch"] == "Sheikh Zayed"
    assert [c["account_no"] for c in data["customers"]] == ["ACC1"]
    # unknown branch → FULL capped list with a warning (model filters), never silently empty
    data2, warnings2 = await gen.fetch_datasets(db_session, ["properties"], "ناموجود")
    assert len(data2["properties"]) == 2 and warnings2
    # branches catalog covers both
    branches = await gen.list_branches(db_session)
    assert "Sheikh Zayed" in branches and "Deira" in branches


# ---------------- endpoint wiring (mocked model) ----------------

async def _seed_model(db_session):
    db_session.add(AIProvider(key="anthropic", display_name="Anthropic", enabled=True,
                              auth_scheme="api_key", base_url="https://api.anthropic.com",
                              api_key="sk-test-xxx"))
    db_session.add(AIModel(model_key="claude-opus-4-8", provider_key="anthropic",
                           display_name="Claude Opus 4.8", enabled=True,
                           capabilities=["text"], priority=1))
    await db_session.commit()


async def test_generate_attachment_endpoint_stores_marked_file(
    client, auth_headers, db_session, monkeypatch,
):
    await _seed_model(db_session)

    async def fake_complete(db, prompt, **kwargs):
        # the wiring must hand the model the DB facts + the user's instruction
        assert "حقایقِ پایگاه‌داده" in prompt
        assert "جدول وضعیت املاک" in prompt
        return {"ok": True, "model": "claude-opus-4-8", "text": json.dumps({
            "kind": "excel", "filename": "وضعیت املاک", "title": "وضعیت املاک رهنی",
            "warnings": ["ارزش کارشناسی در پایگاه‌داده نبود"],
            "sheets": [{"name": "املاک", "columns": ["ردیف", "وضعیت"],
                        "rows": [["۱", "فاقد بیمه‌نامه"]]}],
        }, ensure_ascii=False)}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)

    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "LTRX1", "account_no": "", "kind": "excel",
        "instruction": "جدول وضعیت املاک رهنی را بساز",
        "subject": "استعلام وضعیت املاک",
    })
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["ok"] is True and body["kind"] == "excel"
    assert body["warnings"] == ["ارزش کارشناسی در پایگاه‌داده نبود"]
    att = body["attachment"]
    assert att["original_name"].endswith(".xlsx")
    assert att["ai_generated"] is True

    # the letter's attachment listing must expose the flag (frontend uses it to
    # default-EXCLUDE these from the extraction tool)
    r2 = await client.get("/api/letters/LTRX1/attachments", headers=auth_headers)
    assert r2.status_code == 200
    rows = r2.json()
    assert len(rows) == 1 and rows[0]["ai_generated"] is True

    # and the stored bytes are a real workbook (disk fallback path in tests)
    r3 = await client.get(f"/api/crm/attachments/{att['id']}/download", headers=auth_headers)
    assert r3.status_code == 200
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(r3.content))
    assert wb["املاک"].cell(row=3, column=2).value == "فاقد بیمه‌نامه"


async def test_generate_attachment_branch_wide_two_phase(
    client, auth_headers, db_session, monkeypatch,
):
    """The owner's real failure: «لیست املاک شعبهٔ شیخ زاید» on a general letter
    (no account) returned an empty skeleton. Now the model sees the datasets
    catalog + real branch values, requests need_data, and the second round gets
    the actual branch-filtered rows to build the file from."""
    await _seed_model(db_session)
    await _seed_branch_data(db_session)
    calls = []

    async def fake_complete(db, prompt, **kwargs):
        calls.append(prompt)
        if len(calls) == 1:
            # round 1: catalog + real branch values must be offered
            assert "کاتالوگِ داده‌های سراسری" in prompt
            assert "Sheikh Zayed" in prompt and "properties" in prompt
            return {"ok": True, "model": "m", "text": json.dumps(
                {"need_data": {"datasets": ["properties"], "branch": "Sheikh Zayed"}},
                ensure_ascii=False)}
        # round 2: the fetched, branch-filtered rows are in the prompt
        assert "داده‌های واکشی‌شده" in prompt
        assert "1234/56" in prompt          # Sheikh Zayed property
        assert "9999/99" not in prompt      # Deira property filtered out
        return {"ok": True, "model": "m", "text": json.dumps({
            "kind": "excel", "filename": "املاک شعبه", "title": "املاک رهنی شعبه شیخ زاید",
            "sheets": [{"name": "املاک",
                        "columns": ["حساب", "پلاک ثبتی", "سند رهنی", "بیمه‌نامه", "مدیر حساب"],
                        "rows": [["ACC1", "1234/56", "MD-77", "INS-9", "آقای مدیری"]]}],
        }, ensure_ascii=False)}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)
    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "LTRX9", "account_no": "",   # general letter — the owner's case
        "instruction": "لیستی از املاک شعبه شیخ زاید به همراه وضعیت بیمه‌نامه و شماره اسناد رهنی و پلاک ثبتی و اسامی مدیران حساب‌ها",
    })
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["ok"] is True and len(calls) == 2
    # the produced workbook actually contains the DB rows
    r2 = await client.get(f"/api/crm/attachments/{body['attachment']['id']}/download",
                          headers=auth_headers)
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(r2.content))
    ws = wb["املاک"]
    assert ws.cell(row=3, column=2).value == "1234/56"
    assert ws.cell(row=3, column=5).value == "آقای مدیری"


async def test_extract_refuses_ai_generated_attachment(
    client, auth_headers, db_session, monkeypatch,
):
    """Server-side circular-write guard: an AI-generated attachment must be
    refused by the extraction endpoint unless the override flag is explicit —
    the frontend's default-untick is only UI feedback, the server is the gate."""
    await _seed_model(db_session)

    async def fake_complete(db, prompt, **kwargs):
        return {"ok": True, "model": "m", "text": json.dumps({
            "kind": "excel", "filename": "گزارش", "title": "گزارش",
            "sheets": [{"name": "برگه", "columns": ["الف"], "rows": [["۱"]]}],
        }, ensure_ascii=False)}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)
    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "LTRX3", "instruction": "جدول گزارش بساز",
    })
    assert r.status_code == 200, r.text
    att_id = r.json()["attachment"]["id"]

    # default: refused with a typed error, nothing staged
    r2 = await client.post(f"/api/letter-ai/extract-attachment/{att_id}",
                           headers=auth_headers, json={})
    assert r2.status_code == 200
    assert r2.json()["ok"] is False
    assert r2.json()["error"] == "ai_generated_attachment"

    # explicit override (user deliberately ticked it in the UI) passes the guard
    from app.services import letter_attachment_extract as lax

    async def fake_extract(*args, **kwargs):
        return {"ok": True, "model": "m", "facts": [], "chunk_errors": []}

    async def fake_stage(*args, **kwargs):
        return []

    monkeypatch.setattr(lax, "extract_attachment", fake_extract)
    monkeypatch.setattr(lax, "stage_extraction", fake_stage)
    r3 = await client.post(f"/api/letter-ai/extract-attachment/{att_id}",
                           headers=auth_headers, json={"allow_ai_generated": True})
    assert r3.status_code == 200, r3.text
    assert r3.json()["ok"] is True


async def test_generate_attachment_rejects_bad_spec(client, auth_headers, db_session, monkeypatch):
    await _seed_model(db_session)

    async def fake_complete(db, prompt, **kwargs):
        return {"ok": True, "model": "m", "text": "این اصلاً JSON نیست"}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)
    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "LTRX2", "instruction": "هرچیزی",
    })
    assert r.status_code == 200
    assert r.json()["ok"] is False and r.json()["error"].startswith("bad_spec")


# ---------------- v63: TEMPLATE/SAMPLE-file driven generation ----------------

def test_build_prompt_template_section_and_rules():
    from app.services import letter_attachment_generate as gen
    p = gen.build_prompt({}, {"subject": "س"}, "ستون تاریخ هم اضافه کن",
                         template_text="نام مشتری | شماره حساب | مبلغ", template_name="فرم-اداره.xlsx")
    assert "قالب/نمونهٔ داده‌شده توسط کاربر" in p and "فرم-اداره.xlsx" in p
    assert "نام مشتری | شماره حساب | مبلغ" in p
    assert "عیناً بازتولید" in p
    assert "ستون تاریخ هم اضافه کن" in p           # extras apply ON TOP of the template
    # template-only: the instruction section says build strictly from the template
    p2 = gen.build_prompt({}, {}, "", template_text="ستون‌های الف/ب", template_name="t.pdf")
    assert "بدون شرح — پیوست را دقیقاً مطابقِ قالبِ داده‌شده" in p2
    # SYSTEM rule 9 exists
    assert "قالب/نمونهٔ داده‌شده توسط کاربر" in gen.SYSTEM_PROMPT


async def test_generate_with_template_only(client, auth_headers, db_session, monkeypatch):
    await _seed_model(db_session)
    seen = {}

    async def fake_complete(db, prompt, **kwargs):
        seen["prompt"] = prompt
        return {"ok": True, "model": "claude-opus-4-8", "text": json.dumps({
            "kind": "excel", "filename": "طبق قالب", "title": "فرم اداره",
            "warnings": [],
            "sheets": [{"name": "فرم", "columns": ["نام مشتری", "شماره حساب", "مبلغ"],
                        "rows": [["الف", "۱۲۳", "۱٬۰۰۰"]]}],
        }, ensure_ascii=False)}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)

    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "L-TPL-1", "instruction": "",
        "template_text": "نام مشتری | شماره حساب | مبلغ", "template_name": "فرم-اداره.xlsx",
    })
    assert r.status_code == 200 and r.json()["ok"] is True
    assert "قالب/نمونهٔ داده‌شده" in seen["prompt"] and "فرم-اداره.xlsx" in seen["prompt"]
    # neither instruction nor template → 422 with a clear Persian message
    r2 = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "L-TPL-2", "instruction": "", "template_text": "",
    })
    assert r2.status_code == 422


async def test_template_text_endpoint_reads_xlsx(client, auth_headers):
    import io
    from openpyxl import Workbook
    wb = Workbook(); ws = wb.active; ws.title = "فرم"
    ws.append(["نام مشتری", "شماره حساب", "مبلغ ترهین"])
    buf = io.BytesIO(); wb.save(buf)
    r = await client.post("/api/letter-ai/template-text", headers=auth_headers,
                          files={"file": ("فرم-خالی.xlsx", buf.getvalue(),
                                          "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")})
    assert r.status_code == 200
    body = r.json()
    assert body["ok"] is True and "شماره حساب" in body["text"] and "مبلغ ترهین" in body["text"]


# ---------------- v65: SOURCE/DATA files feed the generation ----------------

def test_build_prompt_source_files_section_and_rule():
    from app.services import letter_attachment_generate as gen
    p = gen.build_prompt({}, {}, "جدول را از این فایل‌ها بساز",
                         source_files=[{"name": "لیست-شعبه.xlsx", "text": "حساب ۱۲۳ | مبلغ ۵۰۰"},
                                       {"name": "گزارش.pdf", "text": "متن گزارش ارزیابی"},
                                       {"name": "empty.txt", "text": "  "}])
    assert "فایل‌های منبعِ داده" in p
    assert "لیست-شعبه.xlsx" in p and "حساب ۱۲۳" in p
    assert "گزارش.pdf" in p and "متن گزارش ارزیابی" in p
    assert "empty.txt" not in p                     # blank-text files are dropped
    # system rules: sources are a legitimate data source; template=SHAPE, sources=DATA
    assert "فایل‌های منبعِ داده" in gen.SYSTEM_PROMPT
    assert "قالب/نمونه» فقط شکلِ خروجی" in gen.SYSTEM_PROMPT
    # template + sources coexist in one prompt, each in its own section
    p2 = gen.build_prompt({}, {}, "", template_text="ستون الف | ستون ب", template_name="فرم.xlsx",
                          source_files=[{"name": "داده.csv", "text": "ردیف ۱"}])
    assert "قالب/نمونهٔ داده‌شده" in p2 and "فایل‌های منبعِ داده" in p2
    assert p2.index("فایل‌های منبعِ داده") < p2.index("قالب/نمونهٔ داده‌شده")


async def test_generate_with_source_files_reaches_model(client, auth_headers, db_session, monkeypatch):
    await _seed_model(db_session)
    seen = {}

    async def fake_complete(db, prompt, **kwargs):
        seen["prompt"] = prompt
        return {"ok": True, "model": "claude-opus-4-8", "text": json.dumps({
            "kind": "excel", "filename": "از منابع", "title": "جدول از فایل‌های منبع",
            "warnings": [],
            "sheets": [{"name": "داده", "columns": ["حساب", "مبلغ"], "rows": [["۱۲۳", "۵۰۰"]]}],
        }, ensure_ascii=False)}

    from app.ai import inference
    monkeypatch.setattr(inference, "complete", fake_complete)

    r = await client.post("/api/letter-ai/generate-attachment", headers=auth_headers, json={
        "letter_id": "L-SRC-1", "instruction": "جدول حساب/مبلغ از فایل منبع بساز",
        "source_files": [{"name": "لیست-شعبه.xlsx", "text": "حساب 123 مبلغ 500"}],
    })
    assert r.status_code == 200 and r.json()["ok"] is True
    assert "فایل‌های منبعِ داده" in seen["prompt"] and "لیست-شعبه.xlsx" in seen["prompt"]


# ---------------- v66: LOGS as need_data datasets ----------------

async def test_fetch_datasets_logs(db_session):
    """Global audit trail + journal/daily-log lines are requestable datasets —
    catalogued, capped, branch-filterable through the row's account."""
    from app.models.audit_log import AuditLog
    from app.models.crm import JournalEntry

    await _seed_branch_data(db_session)
    db_session.add(AuditLog(id="a1", username="u1", action="update", entity_type="letter",
                            account_no="ACC1", detail="ویرایش نامهٔ رسمی"))
    db_session.add(AuditLog(id="a2", username="u2", action="print", entity_type="letter",
                            account_no="ACC2", detail="چاپ نامه"))
    db_session.add(JournalEntry(id="j1", account_no="ACC1", category="letters",
                                item="صدور نامه", status="done", date="2026-07-14",
                                user="u1", notes="نامهٔ ترهین"))
    await db_session.commit()

    # advertised in the catalog + accepted by the need_data parser
    assert "audit_logs" in gen.DATASETS and "journal_entries" in gen.DATASETS
    cat = gen.catalog_text([])
    assert "audit_logs" in cat and "journal_entries" in cat
    need = gen.parse_need_data(json.dumps({"need_data": {"datasets": ["audit_logs", "journal_entries"]}}))
    assert need and set(need["datasets"]) == {"audit_logs", "journal_entries"}

    data, warnings = await gen.fetch_datasets(db_session, ["audit_logs", "journal_entries"], "")
    logs = data["audit_logs"]
    assert {r["account_no"] for r in logs} == {"ACC1", "ACC2"}
    row = next(r for r in logs if r["account_no"] == "ACC1")
    assert row["user"] == "u1" and row["action"] == "update"
    assert row["customer_name"] == "شرکت الف" and "ویرایش" in row["detail"]
    j = data["journal_entries"]
    assert len(j) == 1 and j[0]["item"] == "صدور نامه" and j[0]["customer_name"] == "شرکت الف"

    # branch filter → only the matched account's rows
    data2, _ = await gen.fetch_datasets(db_session, ["audit_logs"], "sheikh zayed")
    assert [r["account_no"] for r in data2["audit_logs"]] == ["ACC1"]


async def test_need_data_logs_filter_full_table_search(db_session):
    """v67: the logs datasets run through the UNLIMITED search (whole table,
    filterable) — need_data carries logs_filter and old rows stay reachable."""
    from datetime import datetime
    from app.models.audit_log import AuditLog

    await _seed_branch_data(db_session)
    db_session.add(AuditLog(id="new1", username="u1", action="update", entity_type="letter",
                            account_no="ACC1", detail="کار جدید",
                            created_at=datetime(2026, 7, 15, 12, 0, 0)))
    db_session.add(AuditLog(id="old1", username="u2", action="print", entity_type="letter",
                            account_no="ACC1", detail="چاپ نامهٔ ترهین بسیار قدیمی",
                            created_at=datetime(2024, 3, 1, 12, 0, 0)))
    await db_session.commit()

    # parse_need_data keeps a sanitized logs_filter (rule 8)
    need = gen.parse_need_data(json.dumps({"need_data": {
        "datasets": ["audit_logs"], "logs_filter": {"text": "ترهین", "junk": "x"}}}))
    assert need and need["logs_filter"] == {"text": "ترهین"}
    assert "logs_filter" in gen.SYSTEM_PROMPT

    # the filter reaches the search: only the OLD matching row comes back
    data, _ = await gen.fetch_datasets(db_session, need["datasets"], "",
                                       logs_filter=need["logs_filter"])
    rows = data["audit_logs"]
    assert len(rows) == 1 and "قدیمی" in rows[0]["detail"]
    assert rows[0]["customer_name"] == "شرکت الف"
