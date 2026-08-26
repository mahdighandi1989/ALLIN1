"""Deep AI extraction from a LETTER's attachments → staged, reviewable changes.

Reuses the battle-tested IMPORT pipeline pieces (they already solved the hard
parts: unreadable files, size limits, mid-run crashes/OOM, rate-limit retries):

  * PDF/images  → ``doc_ingest.pdf_chunks`` (streamed page-chunks) +
                  ``inference.resolve_multimodal`` once + ``send_multimodal``
                  per chunk with a single 429 backoff retry, freeing each
                  chunk's bytes before the next.
  * Excel/CSV   → ``doc_ingest.workbook_to_text`` + ``chunk_text``.
  * Word (.docx)→ the deterministic ``draft_extract`` parser (no AI needed).

The prompt is the import EXTRACTION_PROMPT **plus letter context** (subject,
the letter's own customer, body summary) and a *relationships* section — so the
model extracts EVERYTHING relevant to the letter's subject, for EVERY account
named anywhere in the attachment (not just the primary), thoroughly and without
summarizing, and states how the named parties relate to each other with the
exact reason.

Output is NOT persisted here: extracted facts are staged through the same
``letter_db_extract.stage_db_writes`` gate (add/update/skip + unresolved-never-
guessed) and relationship proposals become reviewable ``link`` items. The user
ticks; ``/apply-db`` writes.
"""
from __future__ import annotations

import asyncio
import json
import logging
from typing import Any, Dict, List, Optional

from sqlalchemy.ext.asyncio import AsyncSession

from app.ai import inference
from app.services import doc_ingest
from app.services import letter_db_extract as dbx

logger = logging.getLogger(__name__)

# Same coverage as the Import page (tiff/bmp scans included) + plain text.
_IMAGE_MIMES = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif",
                "image/tiff", "image/bmp"}
_EXCEL_EXT = (".xlsx", ".xlsm", ".xls", ".csv")
_TEXT_EXT = (".txt", ".text", ".log")
_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# Same chunking thresholds as the import page (proven against OOM/timeouts).
_PDF_SPLIT_BYTES = 4 * 1024 * 1024
_PDF_CHUNK_BYTES = 18 * 1024 * 1024
_PDF_CHUNK_PAGES = 12
_MAX_CHUNKS = 16  # hard cap per attachment — beyond this, use the Import page


def build_prompt(letter_ctx: Dict[str, str], part: str = "") -> str:
    """The import extraction prompt + letter context + relationships section."""
    ctx_lines = [
        "\n\n### LETTER CONTEXT (this file is an attachment of an official letter):",
        f"- Letter subject: {letter_ctx.get('subject') or '-'}",
        f"- Letter's primary customer: {letter_ctx.get('customer_name') or '-'} "
        f"(account {letter_ctx.get('account_no') or '-'})",
    ]
    if letter_ctx.get("body_excerpt"):
        ctx_lines.append(f"- Letter body (for context): {letter_ctx['body_excerpt'][:1500]}")
    ctx_lines.append(
        "Extract EVERYTHING relevant to the letter's subject — completely and precisely, "
        "NO summarizing, NO skipping, NO 'etc.' — for EVERY account/customer named anywhere "
        "in this attachment (not only the primary customer). Attribute each fact to the "
        "correct account; when in doubt about WHO a fact belongs to, leave it out rather "
        "than guessing.\n"
        "ALSO add a top-level \"relationships\" array to the SAME JSON object:\n"
        "[{\"from_account\": \"\", \"from_name\": \"\", \"to_account\": \"\", \"to_name\": \"\", "
        "\"kind\": \"guarantor|co_signer|family|business_partner|letter|other\", \"reason\": \"\"}]\n"
        "Relationship rules:\n"
        "- Record ONLY relationships the document explicitly states or unambiguously shows "
        "(a signature as guarantor, a listed partner, a joint account) — never inferred ones.\n"
        "- \"reason\" must QUOTE or precisely restate the document's own wording (in Persian), "
        "including the document/section it came from — this exact text is stored on both "
        "customer profiles as the recorded justification of the link.\n"
        "- Pick the MOST SPECIFIC kind (a guarantor is \"guarantor\", not \"other\"); use "
        "\"letter\" only for parties connected merely by being named in this correspondence.\n"
        "- Never invent accounts, names, numbers or relationships.\n"
        "ALSO add a top-level \"kb_items\" array to the SAME JSON object for GENERAL, reusable, "
        "NON-customer-specific knowledge this file contains — circulars (بخشنامه) with their exact "
        "numbers/dates, procedures, tariffs, regulatory points, reference tables, aggregate/branch-"
        "level statistics, educational material:\n"
        "[{\"topic\": \"...\", \"category\": \"...\", \"content\": \"...\", \"source_note\": \"...\"}]\n"
        "kb_items rules: content in Persian, complete and precise (keep the document's exact numbers "
        "and dates); source_note cites this file and the section it came from; NEVER put a specific "
        "customer's private data (their name/account/amounts) in kb_items; empty array when the file "
        "has none.\n"
        "If the file has NO per-account customer data at all (an aggregate or branch-level report, a "
        "general/reference letter…), that is NOT an error: return \"customers\": [] and still fill "
        "\"kb_items\" with the general knowledge the file holds."
    )
    return doc_ingest.EXTRACTION_PROMPT + "\n".join(ctx_lines) + (("\n\n" + part) if part else "")


async def extract_attachment(
    db: AsyncSession, *, data: bytes, filename: str, mimetype: str,
    letter_ctx: Dict[str, str], model_id: Optional[int] = None,
) -> Dict[str, Any]:
    """Run the full pipeline over ONE attachment. Returns
    ``{ok, customers, relationships, model, chunk_errors}`` (never persists)."""
    lower = (filename or "").lower()
    is_pdf = mimetype == "application/pdf" or lower.endswith(".pdf")
    customers_merged: Dict[str, dict] = {}
    relationships: List[dict] = []
    kb_items: List[dict] = []
    chunk_errors: List[str] = []
    model_name = None

    def _fold(parsed: dict) -> None:
        for c in (parsed.get("customers") or []):
            a = doc_ingest._acc_of(c)
            if not a:
                continue
            if a in customers_merged:
                doc_ingest.merge_customer(customers_merged[a], c)
            else:
                customers_merged[a] = c
        for r in (parsed.get("relationships") or []):
            if isinstance(r, dict):
                relationships.append(r)
        for k in (parsed.get("kb_items") or []):
            if isinstance(k, dict):
                kb_items.append(k)

    if lower.endswith(_TEXT_EXT) or mimetype == "text/plain":
        # Plain-text attachment: decode tolerantly, chunk, extract as text.
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("cp1256", errors="replace")  # Arabic/Persian legacy encoding
        if not text.strip():
            return {"ok": False, "error": "فایل متنی خالی است."}
        chunks = doc_ingest.chunk_text(text, 100000)[:_MAX_CHUNKS]
        for ci, ch in enumerate(chunks):
            prompt = build_prompt(
                letter_ctx,
                f"The content below is a PLAIN-TEXT document (part {ci+1} of {len(chunks)}). "
                "There are no page images — ignore the 'documents' array.\n\nTEXT CONTENT:\n" + ch,
            )
            res = await inference.complete(db, prompt, task="document_extraction",
                                           model_id=model_id, max_tokens=8000, timeout=180.0)
            if not res.get("ok"):
                if res.get("error") == "no_model":
                    return {"ok": False, "error": "no_model"}
                chunk_errors.append(str(res.get("error")))
                continue
            model_name = res.get("model")
            _fold(doc_ingest.parse_model_json(res.get("text", "")))
    elif lower.endswith(_EXCEL_EXT) or mimetype in ("text/csv",):
        try:
            table_text = doc_ingest.workbook_to_text(data, filename)
        except Exception as exc:
            return {"ok": False, "error": f"جدول قابلِ خواندن نبود: {exc}"}
        # DETERMINISTIC FIRST for tables: an account-column sheet maps instantly
        # and this endpoint runs INLINE in the HTTP request — a long model call
        # here blew the platform gateway timeout and surfaced as a bare 502.
        # The model pass stays available for sheets the parser cannot map.
        for c in doc_ingest.table_fallback_customers(table_text):
            acc = doc_ingest._acc_of(c)
            if acc:
                customers_merged[acc] = c
        if customers_merged:
            model_name = "استخراج قطعی جدول"
        else:
            chunks = doc_ingest.chunk_text(table_text, 100000)[:_MAX_CHUNKS]
            for ci, ch in enumerate(chunks):
                prompt = build_prompt(
                    letter_ctx,
                    f"The content below is a SPREADSHEET/TABLE (part {ci+1} of {len(chunks)}). "
                    "Extract EVERY row, attributing each to its account.\n\nTABLE CONTENT:\n" + ch,
                )
                res = await inference.complete(db, prompt, task="document_extraction",
                                               model_id=model_id, max_tokens=8000, timeout=75.0)
                if not res.get("ok"):
                    if res.get("error") == "no_model":
                        return {"ok": False, "error": "no_model"}
                    chunk_errors.append(str(res.get("error")))
                    continue
                model_name = res.get("model")
                _fold(doc_ingest.parse_model_json(res.get("text", "")))
    elif mimetype == _DOCX_MIME or lower.endswith(".docx"):
        from app.services.draft_extract import extract_from_docx
        try:
            dx = extract_from_docx(data)
        except Exception as exc:
            return {"ok": False, "error": f"فایل Word قابلِ خواندن نبود: {exc}"}
        o, pf = dx.get("offer", {}), dx.get("profile", {})
        if dx.get("account_no"):
            customers_merged[dx["account_no"]] = {
                "account_no": dx.get("account_no"), "name": o.get("CompanyName"),
                "fields": {**pf}, "guarantors": dx.get("guarantors", []),
            }
        model_name = "Word draft parser"
    elif is_pdf or mimetype in _IMAGE_MIMES:
        # Streamed page-chunks so the full set is never in memory (import lesson).
        if is_pdf and len(data) > _PDF_SPLIT_BYTES:
            try:
                chunk_iter = doc_ingest.pdf_chunks(data, max_bytes=_PDF_CHUNK_BYTES,
                                                   max_pages=_PDF_CHUNK_PAGES)
            except Exception as exc:
                return {"ok": False, "error": f"فایلِ بزرگ قابلِ تقسیم نبود: {exc}"}
        else:
            chunk_iter = iter([(0, data)])
        rr = await inference.resolve_multimodal(db, [{"mimetype": mimetype}], model_id=model_id)
        if not rr.get("ok"):
            return {"ok": False, "error": rr.get("error"),
                    "suggestions": rr.get("suggestions", []), "model": rr.get("model")}
        resolved = rr["resolved"]
        model_name = resolved.display_name
        prompt = build_prompt(letter_ctx)
        n = 0
        for _start, pbytes in chunk_iter:
            n += 1
            if n > _MAX_CHUNKS:
                chunk_errors.append("سقفِ قطعات پر شد — برای فایل‌های خیلی بزرگ از صفحهٔ Import استفاده کن")
                break
            res = await inference.send_multimodal(
                resolved, prompt, [{"filename": filename, "mimetype": mimetype, "data": pbytes}],
                max_tokens=8000)
            if not res.get("ok") and "429" in str(res.get("error", "")):
                await asyncio.sleep(3)  # single backoff retry on rate-limit (import lesson)
                res = await inference.send_multimodal(
                    resolved, prompt, [{"filename": filename, "mimetype": mimetype, "data": pbytes}],
                    max_tokens=8000)
            pbytes = None  # free this chunk before the next (import lesson)
            if not res.get("ok"):
                chunk_errors.append(str(res.get("error")))
                continue
            _fold(doc_ingest.parse_model_json(res.get("text", "")))
    else:
        return {"ok": False, "error": "فرمتِ این پیوست پشتیبانی نمی‌شود (PDF/تصویر/Excel/CSV/Word)."}

    if not customers_merged and not relationships and not kb_items:
        # Every chunk succeeded but nothing came out ⇒ the file simply holds no
        # account-keyed data (aggregate/branch reports, general letters). That
        # is a fact worth reporting clearly, NOT a scary "model error".
        return {"ok": False,
                "error": (chunk_errors[0] if chunk_errors else "no_account_data"),
                "chunk_errors": chunk_errors}
    return {"ok": True, "customers": list(customers_merged.values()),
            "relationships": relationships, "kb_items": kb_items,
            "model": model_name, "chunk_errors": chunk_errors}


def _flatten_fields(cust: dict) -> Dict[str, str]:
    """One flat {key: value} per extracted customer (fields + top-level extras)."""
    out: Dict[str, str] = {}
    fields = cust.get("fields") if isinstance(cust.get("fields"), dict) else {}
    for k, v in fields.items():
        if v not in (None, "", "-") and not isinstance(v, (list, dict)):
            out[str(k)] = str(v)
    for k in ("account_type", "branch", "business_type"):
        v = cust.get(k)
        if v not in (None, "", "-"):
            out[k] = str(v)
    return out


async def stage_extraction(
    db: AsyncSession, extraction: Dict[str, Any], *, primary_account: str,
    primary_name: str, source_ref: str,
) -> List[Dict[str, Any]]:
    """Turn one attachment's extraction into reviewable items:
    db_write facts (through the SAME stage_db_writes gate) + link proposals."""
    raw_writes: List[dict] = []
    for cust in extraction.get("customers") or []:
        acc = doc_ingest._acc_of(cust) or ""
        name = str(cust.get("name") or "").strip()
        for k, v in _flatten_fields(cust).items():
            raw_writes.append({
                "account_no": acc, "customer_name": name,
                "key": k, "value": v, "title": "", "detail": f"از پیوست: {source_ref}",
            })
        # guarantors named in the attachment → relationship proposals
        for g in cust.get("guarantors") or []:
            gname = str((g or {}).get("guarantor_name") or (g or {}).get("name") or "").strip()
            gacc = str((g or {}).get("guarantor_account") or (g or {}).get("account") or "").strip()
            if gacc and acc:
                extraction.setdefault("relationships", []).append({
                    "from_account": gacc, "from_name": gname, "to_account": acc,
                    "to_name": name, "kind": "guarantor",
                    "reason": f"ضامن طبق پیوستِ نامه ({source_ref})",
                })

    from app.services.letter_assistant import _norm_key
    for w in raw_writes:
        w["key"] = _norm_key(str(w["key"]))
    raw_writes = [w for w in raw_writes if w["key"] and w["value"]]

    staged = await dbx.stage_db_writes(db, primary_account, primary_name, raw_writes)

    # Relationship/link proposals (deduped by pair+kind) — reviewable, applied later.
    seen: set = set()
    for i, r in enumerate(extraction.get("relationships") or []):
        fa_, ta = str(r.get("from_account") or "").strip(), str(r.get("to_account") or "").strip()
        kind = str(r.get("kind") or "other").strip()
        reason = " ".join(str(r.get("reason") or "").split())
        if not fa_ or not ta or fa_ == ta or not reason:
            continue
        key = (tuple(sorted((fa_, ta))), kind)
        if key in seen:
            continue
        seen.add(key)
        staged.append({
            "id": f"l{i}", "op": "link", "category": "db_extract", "field": kind,
            "account_no": fa_, "customer_name": str(r.get("from_name") or ""),
            "related_account": ta, "related_name": str(r.get("to_name") or ""),
            "kind": kind, "reason": reason, "value": reason,
            "title": f"لینکِ «{kind}»: {fa_} ↔ {ta}",
            "detail": reason, "severity": "medium", "applicable": True,
        })

    # GENERAL/reference knowledge from the file → reviewable kb_write items, so
    # a file with no account-keyed rows (aggregate report, circular, general
    # letter) still yields something valuable instead of «nothing extracted».
    seen_kb: set = set()
    kb_n = 0
    for k in extraction.get("kb_items") or []:
        if not isinstance(k, dict):
            continue
        topic = str(k.get("topic") or "").strip()
        content = str(k.get("content") or "").strip()
        if len(topic) < 2 or len(content) < 10:
            continue
        dd = (topic.casefold(), content[:120].casefold())
        if dd in seen_kb:
            continue
        seen_kb.add(dd)
        kb_n += 1
        staged.append({
            "id": f"kb{kb_n}", "op": "kb_write", "category": "db_extract", "field": "",
            "severity": "low", "applicable": True,
            "title": f"پایگاه دانش: {topic[:60]}",
            "detail": f"محتوای عمومی/مرجع از پیوستِ «{source_ref}»",
            "topic": topic[:300],
            "kb_category": str(k.get("category") or "").strip()[:120],
            "content": content[:8000],
            "source_note": (str(k.get("source_note") or "").strip()
                            or f"پیوستِ نامه: {source_ref}")[:400],
        })
        if kb_n >= 20:
            break
    return staged


# ---------------------------------------------------------------------------
# Attachment TEXT (v61): a bounded, non-persisting transcription of ONE
# attachment so the letter assistant can cross-check the letter against its
# attachments and run the full consistency pass. Deterministic wherever
# possible (txt/Excel/CSV/docx); PDF/image go through ONE bounded multimodal
# transcription call (the frontend runs attachments sequentially, one HTTP
# request each, so every call stays under the gateway limit).
# ---------------------------------------------------------------------------
# v114 — the 20k cap silently amputated long PDFs (a 30-page source kept only
# ~7 pages and the generated attachment came out badly incomplete — owner's
# report). PDFs are now transcribed CHUNK BY CHUNK (every page reaches the
# model) and the per-attachment budget is big enough for long documents; the
# attachment-generation prompt applies its own total budget on top.
_TEXT_CAP = 120000         # chars returned per attachment (prompt budget)
_TRANSCRIBE_MAX_BYTES = 40 * 1024 * 1024
_PDF_CHUNK_BYTES = 4 * 1024 * 1024   # per-transcription-call page-chunk bounds
_PDF_CHUNK_PAGES = 8


async def attachment_text(
    db: AsyncSession, *, data: bytes, filename: str, mimetype: str,
    model_id: Optional[int] = None,
) -> Dict[str, Any]:
    """Return ``{ok, text, model?}`` — the attachment's readable content.
    NEVER writes anything; errors come back as {ok: False, error}."""
    lower = (filename or "").lower()
    is_pdf = mimetype == "application/pdf" or lower.endswith(".pdf")

    if lower.endswith(_TEXT_EXT) or mimetype == "text/plain":
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("cp1256", errors="replace")
        return {"ok": True, "text": text[:_TEXT_CAP], "truncated": len(text) > _TEXT_CAP}

    if lower.endswith(_EXCEL_EXT) or mimetype in ("text/csv",):
        try:
            text = doc_ingest.workbook_to_text(data, filename)
            return {"ok": True, "text": text[:_TEXT_CAP], "truncated": len(text) > _TEXT_CAP}
        except Exception as exc:
            return {"ok": False, "error": f"جدول قابلِ خواندن نبود: {exc}"}

    if mimetype == _DOCX_MIME or lower.endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if (p.text or "").strip()]
            for tb in doc.tables:
                for row in tb.rows:
                    parts.append(" | ".join((c.text or "").strip() for c in row.cells))
            text = "\n".join(parts)
            return {"ok": True, "text": text[:_TEXT_CAP], "truncated": len(text) > _TEXT_CAP}
        except Exception as exc:
            return {"ok": False, "error": f"فایل Word قابلِ خواندن نبود: {exc}"}

    if is_pdf or (mimetype or "").startswith("image/"):
        if len(data) > _TRANSCRIBE_MAX_BYTES:
            return {"ok": False, "error": "فایل برای رونویسی بزرگ است (سقف ۴۰MB) — آن را کوچک‌تر کن."}
        rr = await inference.resolve_multimodal(db, [{"mimetype": mimetype}], model_id=model_id)
        if not rr.get("ok"):
            return {"ok": False, "error": rr.get("error")}
        prompt = (
            "Transcribe ALL the text of the attached document faithfully, page by page, "
            "top to bottom — every heading, paragraph, table (as pipe-separated rows), "
            "reference number, date and amount, exactly as printed. Persian stays Persian, "
            "English stays English. Output ONLY the transcription, no commentary."
        )
        # v114 — a single call's OUTPUT-token cap silently cut long PDFs; split
        # into page-chunks so EVERY page is transcribed, with one retry per
        # chunk on a transient failure. A failed chunk is reported loudly,
        # never skipped in silence.
        if is_pdf:
            try:
                chunks = list(doc_ingest.pdf_chunks(data, max_bytes=_PDF_CHUNK_BYTES,
                                                    max_pages=_PDF_CHUNK_PAGES))
            except Exception:
                chunks = [(0, data)]
        else:
            chunks = [(0, data)]
        import asyncio as _aio
        # send_multimodal is DB-free by design so page-chunks can run
        # CONCURRENTLY (bounded) — this whole transcription happens inside one
        # HTTP request, and a big PDF done serially would blow the gateway
        # deadline that a single bounded call used to stay under.
        sem = _aio.Semaphore(3)

        async def _one(start: int, cbytes: bytes):
            async with sem:
                res = await inference.send_multimodal(
                    rr["resolved"], prompt,
                    [{"filename": filename, "mimetype": mimetype, "data": cbytes}], max_tokens=8000)
                err0 = str(res.get("error") or "")
                if not res.get("ok") and ("timed out" in err0 or "connection failed" in err0 or "429" in err0):
                    await _aio.sleep(3)
                    res = await inference.send_multimodal(
                        rr["resolved"], prompt,
                        [{"filename": filename, "mimetype": mimetype, "data": cbytes}], max_tokens=8000)
            return start, res

        results = await _aio.gather(*(_one(s, b) for s, b in chunks))
        parts: list = []
        failed: list = []
        for start, res in results:  # gather keeps input order ⇒ pages stay in order
            if res.get("ok"):
                parts.append((res.get("text") or "").strip())
            else:
                failed.append(start or 1)
                parts.append(f"[⚠ رونویسیِ بخشِ شروع‌شده از صفحهٔ {start or 1} ناموفق بود: {res.get('error')}]")
        if failed and len(failed) == len(chunks):
            return {"ok": False, "error": f"رونویسی ناموفق بود: {parts[0] if parts else ''}"}
        text = "\n\n".join(p for p in parts if p)
        return {"ok": True, "text": text[:_TEXT_CAP],
                "truncated": len(text) > _TEXT_CAP,
                "failed_parts": failed,
                "model": rr["resolved"].display_name}

    return {"ok": False, "error": "فرمتِ این پیوست پشتیبانی نمی‌شود (PDF/تصویر/Excel/CSV/Word/متن)."}
