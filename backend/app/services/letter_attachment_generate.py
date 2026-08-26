"""Generate a REAL file attachment for a letter from the owner's instruction.

Review-first, server-is-ground-truth: the model only ever proposes a STRICT
JSON *spec* (kind/filename/title + sheets for Excel or paragraphs for Word);
a deterministic validator clamps everything, and the actual file is rendered
SERVER-SIDE (openpyxl / python-docx) and stored through the very same
Drive+disk+DB path as manual uploads. The stored record carries an
``AI_GENERATED`` marker in its notes so the attachment-extraction tool can
exclude these by default — their data came out of the database in the first
place (feeding them back in would be a circular write).

Data discipline mirrors the letter assistant: values come ONLY from the
database facts or the instruction itself; anything unavailable stays EMPTY and
is reported in ``warnings`` — never invented.
"""
from __future__ import annotations

import io
import json
import re
from typing import Any, Dict, List, Tuple

AI_GENERATED_MARK = "AI_GENERATED"

MAX_SHEETS = 3
MAX_ROWS = 500
MAX_COLS = 30
MAX_PARAGRAPHS = 150
MAX_CELL = 500
MAX_DATASET_ROWS = 300

# v114 — the owner's «the attachment built from my source file was badly
# incomplete»: build_prompt used to hard-slice every source file (and the
# template) at 20k chars, silently amputating anything past ~10 pages. The
# caps now fit a whole transcription (letter_attachment_extract caps its text
# at 120k) and any cut is EXPLICIT — an inline marker inside the prompt plus a
# Persian warning surfaced in the API reply (prompt_size_warnings). A very
# large total prompt may still exceed a small model's context window; that
# fails loudly with the provider's error, never silently with missing rows.
SRC_FILE_CAP = 120_000
SRC_TOTAL_CAP = 360_000
TEMPLATE_CAP = 120_000
MAX_SOURCE_FILES = 8

# Cross-customer datasets the model may REQUEST via the need_data protocol —
# a single-account letter (or a general one) can still build branch-wide /
# bank-wide lists: the model names the dataset(s) + an exact branch value, the
# server runs the deterministic capped query, and a second round produces the
# spec from those rows. The model never queries the DB itself.
DATASETS: Dict[str, str] = {
    "properties": ("املاک رهنی: شمارهٔ حساب، نام مشتری، شعبه، مدیرِ حساب، پلاک ثبتی، شمارهٔ سند رهنی، "
                   "شهر، آدرس، کد پستی، نوع، مالک/راهن + کد ملی، متراژ زمین/زیربنا، سن بنا، منطقه، "
                   "ارزیابی + تاریخ آخرین ارزیابی، تاریخ/مبلغ ترهین، ملاحظات، و بیمه‌نامهٔ کامل: "
                   "شماره، کد رایانه، تاریخ صدور/انقضا، بیمه‌گذار، مورد بیمه، شرح دقیق فعالیت شغلی، "
                   "مجموع سرمایهٔ تحت پوشش، واحد کاری صدور"),
    "customers": "مشتریان: شمارهٔ حساب، نام، شعبه، مدیرِ حساب",
    "facilities": "تسهیلات: حساب، نام مشتری، شعبه، مدیرِ حساب، نوع/نام تسهیلات، مبلغ، مانده، نرخ، تاریخ‌های شروع/پایان/انقضا",
    "securities": "تضامین/وثایق چندساله: سال، شعبه، حساب، نام مشتری، FD، ضامن، چک‌ها، مبلغ چک، شمارهٔ ملک، مبلغ ترهین، ملاحظات",
    "fixed_deposits": "سپرده‌های ثابت: حساب، نام مشتری، شمارهٔ سپرده، مبلغ، ارز، تاریخ افتتاح/سررسید، نرخ",
    "guarantors": "ضامن‌ها: حساب، نام مشتری، نام ضامن، حسابِ ضامن، شمارهٔ چک، مبلغ چک، بانک",
    "audit_logs": "لاگِ کلیِ سیستم (Audit Log — جدیدترین اول): تاریخ/زمان، کاربر، عملیات، نوعِ موجودیت، شمارهٔ حساب، شرحِ کار — جستجو روی «کلِ» لاگ‌هاست و با logs_filter (متن/کاربر/عملیات/حساب/بازهٔ تاریخ) محدودش کن",
    "journal_entries": "لاگِ کارها و ثبت‌های روزانه (جدیدترین اول): حساب، نام مشتری، شعبه، دسته، مورد، وضعیت، تاریخ/زمان، کاربر، یادداشت — جستجو روی «کلِ» لاگ‌هاست و logs_filter می‌پذیرد",
}

SYSTEM_PROMPT = (
    "تو سازندهٔ «پیوستِ رسمیِ» یک نامهٔ بانکی هستی. بر اساسِ دستورِ کاربر، زمینهٔ نامه و "
    "«حقایقِ پایگاه‌داده» فقط و فقط یک شیءِ JSON برگردان — بدونِ متنِ اضافه، بدونِ markdown.\n"
    "قواعدِ الزامی:\n"
    "1) دادهٔ واقعی فقط از «حقایقِ پایگاه‌داده»، «فایل‌های منبعِ» داده‌شده توسط کاربر (اگر در پیام "
    "هست) یا خودِ دستورِ کاربر می‌آید؛ هرگز عدد/تاریخ/نام/مبلغ نساز. اگر قلمی خواسته شده و در "
    "هیچ‌کدام موجود نیست، خانه/بخش را خالی بگذار و دلیل را در warnings بنویس.\n"
    "2) لحن و نگارش باید هم‌سطحِ نامهٔ رسمیِ بانکی و هماهنگ با متنِ نامه باشد (رسمی، سوم‌شخص، بدونِ اغراق).\n"
    "3) kind را خودت هوشمندانه انتخاب کن مگر کاربر صریحاً گفته باشد: جدولی/عددی ⇒ \"excel\"، "
    "متنی/توضیحی ⇒ \"word\".\n"
    "4) ساختارِ خروجی:\n"
    "{\"kind\": \"excel\"|\"word\", \"filename\": \"نامِ کوتاهِ فارسیِ فایل بدون پسوند\", "
    "\"title\": \"عنوانِ سند\", \"warnings\": [\"...\"],\n"
    " \"sheets\": [{\"name\": \"...\", \"columns\": [\"...\"], \"rows\": [[\"...\"]]}]   ← فقط برای excel\n"
    " \"paragraphs\": [{\"text\": \"...\", \"heading\": true|false, \"bold\": true|false, "
    "\"align\": \"right\"|\"center\"|\"justify\"}]   ← فقط برای word }\n"
    f"5) سقف‌ها: {MAX_SHEETS} شیت، {MAX_ROWS} ردیف، {MAX_COLS} ستون، {MAX_PARAGRAPHS} پاراگراف.\n"
    "6) سرستون‌ها فارسی و روشن؛ اعدادِ مبلغ با جداکنندهٔ هزارگان؛ ستونِ آخرِ هر شیت را اگر مفید است "
    "«ملاحظات» بگذار.\n"
    "7) filename: فارسی و توصیفِ دقیقِ محتوا — نوعِ گزارش/موضوع + دوره/سال اگر دارد (مثلاً «گزارش تسهیلات پرداختی شعب 2024 تا 2026»)؛ نه عبارتِ عمومی مثل «پیوست» یا «گزارش». بدونِ / \\ : * ? \" < > |. شمارهٔ حساب را خودِ سرور به نام اضافه می‌کند.\n"
    "8) اگر برای انجامِ دستور به داده‌هایی فراتر از «حقایقِ پایگاه‌داده» نیاز داری (فهرستِ چندمشتری، "
    "به تفکیکِ شعبه، یا سراسری) و «کاتالوگِ داده‌های سراسری» در پیام هست، به‌جای spec فقط این JSON را "
    "برگردان تا داده برایت واکشی شود:\n"
    "{\"need_data\": {\"datasets\": [\"<از نام‌های کاتالوگ>\"], \"branch\": \"<دقیقاً یکی از مقادیرِ "
    "فهرست‌شدهٔ شعبه یا رشتهٔ خالی برای همه>\", \"logs_filter\": {\"text\": \"\", \"user\": \"\", "
    "\"action\": \"\", \"account_no\": \"\", \"date_from\": \"YYYY-MM-DD\", \"date_to\": \"YYYY-MM-DD\"}}}\n"
    "logs_filter فقط برای datasetهای لاگ است و اختیاری؛ جستجوی لاگ روی «کلِ» جدول‌ها اجرا می‌شود "
    "(بدونِ محدودیتِ قدمت) و اگر یافته‌ها از سقفِ ارسال بیشتر شود، شمارِ واقعی در warnings می‌آید — "
    "فیلتر را دقیق انتخاب کن.\n"
    "این فرصت فقط یک بار است: در نوبتِ بعد «داده‌های واکشی‌شده» را می‌گیری و باید spec نهایی را بدهی. "
    "نامِ شعبه در دستورِ کاربر ممکن است فارسی و در پایگاه‌داده لاتین باشد — خودت معادلِ درست را از فهرست "
    "انتخاب کن. هرگز به‌جای درخواستِ داده، جدولِ خالی یا دادهٔ ساختگی نده.\n"
    "9) اگر بخشِ «قالب/نمونهٔ داده‌شده توسط کاربر» در پیام هست، آن فرمت الزامی است: ساختار، سرستون‌ها، "
    "ترتیبِ ستون‌ها، عنوان‌ها و بخش‌بندی را «عیناً» همان‌طور بازتولید کن (این قالب را اداره/مرجعِ دیگری "
    "خواسته و باید بشناسدش) و فقط خانه‌ها/بخش‌های داده را از «حقایقِ پایگاه‌داده» پر کن. اگر دستورِ "
    "کاربر چیزی فراتر از قالب خواسته (ستون/بخش/فیلترِ اضافه)، آن را روی همان قالب اضافه کن — قالب را "
    "نشکن. اگر دستوری نیست، خواسته را از خودِ قالب و زمینهٔ نامه بفهم. kind را از شکلِ قالب بگیر "
    "(جدول ⇒ excel، فرمِ متنی ⇒ word) مگر کاربر صریحاً گفته باشد.\n"
    "10) اگر بخشِ «فایل‌های منبعِ داده» در پیام هست، محتوای آن‌ها منبعِ دادهٔ مجاز و هم‌ارزِ "
    "«حقایقِ پایگاه‌داده» است — پیوست را از همین اطلاعات (به‌علاوهٔ حقایقِ DB اگر لازم شد) بساز، "
    "طبقِ دستورِ کاربر. تفکیکِ نقش‌ها را قاطی نکن: «قالب/نمونه» فقط شکلِ خروجی را تعیین می‌کند و "
    "«فایل‌های منبع» فقط داده می‌دهند؛ اگر هر دو هستند، داده‌های فایل‌های منبع (و DB) را دقیقاً در "
    "همان قالب بنشان و خواسته‌های اضافهٔ دستور را هم حتماً اجرا کن. اگر داده‌ای بینِ فایل‌های منبع "
    "و پایگاه‌داده ناسازگار بود، مقدارِ فایلِ منبع را بگذار و مغایرت را در warnings اعلام کن."
)

_JSON_RE = re.compile(r"\{.*\}", re.DOTALL)


def fit_sources(source_files: List[Dict[str, str]] | None) -> Tuple[List[Tuple[str, str, int]], List[str]]:
    """Apply the v114 source budgets ONCE — build_prompt renders exactly what
    this returns and the endpoint surfaces the same warnings, so the prompt and
    the API reply can never disagree about what was cut.

    Returns ``([(name, fitted_text, dropped_chars)], warnings)``."""
    srcs = [s for s in (source_files or []) if isinstance(s, dict) and (s.get("text") or "").strip()]
    warns: List[str] = []
    if len(srcs) > MAX_SOURCE_FILES:
        warns.append(
            f"فقط {MAX_SOURCE_FILES} فایلِ منبعِ نخست به مدل داده شد ({len(srcs)} فایل ارسال شده بود)"
        )
    out: List[Tuple[str, str, int]] = []
    budget = SRC_TOTAL_CAP
    for s in srcs[:MAX_SOURCE_FILES]:
        nm = str(s.get("name") or "فایل منبع")[:120]
        txt = str(s.get("text") or "").strip()
        cut = txt[: min(SRC_FILE_CAP, max(0, budget))]
        budget -= len(cut)
        dropped = len(txt) - len(cut)
        if dropped:
            warns.append(
                f"فایلِ منبع «{nm}» کامل به مدل نرسید ({len(cut):,} از {len(txt):,} نویسه) — "
                "برای منابعِ خیلی حجیم مسیرِ مطمئن‌تر: اول استخراج در صفحهٔ ایمپورت، بعد ساختِ پیوست از داده‌های پایگاه"
            )
        out.append((nm, cut, dropped))
    return out, warns


def prompt_size_warnings(source_files: List[Dict[str, str]] | None, template_text: str = "") -> List[str]:
    """Persian warnings for ANY prompt-side truncation (v114 — a cut source or
    template must never be silent in the API reply)."""
    warns = fit_sources(source_files)[1]
    tt = (template_text or "").strip()
    if len(tt) > TEMPLATE_CAP:
        warns.append(f"فایلِ قالب کامل به مدل نرسید ({TEMPLATE_CAP:,} از {len(tt):,} نویسه)")
    return warns


def build_prompt(
    facts: Dict[str, Any],
    letter_ctx: Dict[str, str],
    instruction: str,
    *,
    catalog: str = "",
    fetched: Dict[str, Any] | None = None,
    template_text: str = "",
    template_name: str = "",
    source_files: List[Dict[str, str]] | None = None,
) -> str:
    parts: List[str] = []
    parts.append("### زمینهٔ نامه (برای لحن و موضوع):")
    for k, label in (("subject", "موضوع"), ("recipient", "گیرنده"), ("body_excerpt", "گزیدهٔ متن")):
        v = (letter_ctx.get(k) or "").strip()
        if v:
            parts.append(f"- {label}: {v[:1200]}")
    parts.append("\n### حقایقِ پایگاه‌داده (تنها منبعِ مجازِ داده):")
    parts.append(json.dumps(facts, ensure_ascii=False, indent=1) if facts else "(بدون رکوردِ تک‌مشتری — اگر دستور به داده‌های سراسری نیاز دارد از need_data استفاده کن)")
    if catalog:
        parts.append("\n### کاتالوگِ داده‌های سراسری (در صورتِ نیاز با need_data درخواست بده — قاعدهٔ ۸):")
        parts.append(catalog)
    if fetched is not None:
        parts.append("\n### داده‌های واکشی‌شده از پایگاه‌داده (پاسخِ need_data تو — تنها منبعِ مجازِ داده):")
        parts.append(json.dumps(fetched, ensure_ascii=False, separators=(",", ":")))
        parts.append("دیگر need_data مجاز نیست؛ همین حالا spec نهایی را از همین داده‌ها بساز. "
                     "اگر پس از فیلترِ درست هیچ ردیفی نماند، جدول را خالی بده و دلیل را در warnings بنویس.")
    srcs, _src_warns = fit_sources(source_files)
    if srcs:
        parts.append("\n### فایل‌های منبعِ داده (قاعدهٔ ۱۰ — دادهٔ پیوست از این‌ها هم می‌آید):")
        for nm, txt, dropped in srcs:
            marker = (
                f"\n[⚠ ادامهٔ این فایل ({dropped:,} نویسه) به سقفِ حجم نرسید — ناقص‌بودنِ داده را در warnings اعلام کن]"
                if dropped else ""
            )
            parts.append(f"[فایلِ منبع: {nm}]\n{txt}{marker}")
    if (template_text or "").strip():
        parts.append(f"\n### قالب/نمونهٔ داده‌شده توسط کاربر ({(template_name or 'فایل نمونه')[:120]}) — فرمتِ الزامیِ خروجی (قاعدهٔ ۹):")
        tt = template_text.strip()
        parts.append(tt[:TEMPLATE_CAP] + (
            "\n[⚠ ادامهٔ قالب به سقفِ حجم نرسید — ناقص‌بودنِ قالب را در warnings اعلام کن]"
            if len(tt) > TEMPLATE_CAP else ""
        ))
        parts.append(
            "قواعدِ قالب: ساختار/سرستون‌ها/ترتیب/عنوان‌ها را عیناً بازتولید کن؛ فقط داده‌ها را از "
            "«حقایقِ پایگاه‌داده» پر کن؛ خواسته‌های اضافهٔ دستورِ کاربر را روی همین قالب اعمال کن؛ "
            "سطرهای نمونه/مثالِ داخلِ قالب داده نیستند — جای آن‌ها دادهٔ واقعی بنشیند."
        )
    parts.append("\n### دستورِ کاربر (پیوستی که باید ساخته شود):")
    inst = instruction.strip()[:3000]
    parts.append(inst or "(بدون شرح — پیوست را دقیقاً مطابقِ قالبِ داده‌شده و زمینهٔ نامه بساز و از پایگاه‌داده پر کن)")
    parts.append("\nحالا فقط JSON را برگردان.")
    return "\n".join(parts)


def catalog_text(branches: List[str]) -> str:
    lines = [f"- {name}: {desc}" for name, desc in DATASETS.items()]
    lines.append(
        "مقادیرِ موجودِ «شعبه» در پایگاه‌داده: "
        + (", ".join(f"«{b}»" for b in branches) if branches else "(هیچ شعبه‌ای ثبت نشده)")
    )
    return "\n".join(lines)


def parse_need_data(raw_text: str) -> Dict[str, Any] | None:
    """Detect a need_data request in the model's reply (None when it's a spec)."""
    m = _JSON_RE.search(raw_text or "")
    if not m:
        return None
    try:
        data = json.loads(m.group(0))
    except Exception:  # noqa: BLE001
        return None
    need = data.get("need_data") if isinstance(data, dict) else None
    if not isinstance(need, dict):
        return None
    datasets = [str(d).strip() for d in (need.get("datasets") or []) if str(d).strip() in DATASETS]
    if not datasets:
        return None
    out = {"datasets": datasets[:4], "branch": str(need.get("branch") or "").strip()[:100]}
    if isinstance(need.get("logs_filter"), dict):
        from app.services.log_search import sanitize_query

        lf = sanitize_query(need["logs_filter"])
        lf.pop("scope", None)  # scope follows the requested dataset names
        if lf:
            out["logs_filter"] = lf
    return out


def _s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def _cap(label: str, items: List[Dict[str, str]], warnings: List[str]) -> List[Dict[str, str]]:
    if len(items) > MAX_DATASET_ROWS:
        warnings.append(f"فهرستِ {label} به {MAX_DATASET_ROWS} ردیفِ نخست محدود شد ({len(items)} ردیف موجود است)")
        return items[:MAX_DATASET_ROWS]
    return items


async def list_branches(db) -> List[str]:
    """Distinct branch values (customers + facilities) for the prompt catalog."""
    from sqlalchemy import select

    from app.models.customer import Customer
    from app.models.facility import Facility

    vals: set = set()
    for col, flt in (
        (Customer.branch, Customer.is_deleted == False),  # noqa: E712
        (Facility.branch, Facility.is_deleted == False),  # noqa: E712
    ):
        try:
            rows = (await db.execute(select(col).where(flt).distinct())).scalars().all()
            vals.update(str(v).strip() for v in rows if v and str(v).strip())
        except Exception:  # noqa: BLE001 - a missing table must not kill generation
            continue
    return sorted(vals)[:60]


async def fetch_datasets(db, datasets: List[str], branch: str = "",
                         logs_filter: Dict[str, str] | None = None) -> Tuple[Dict[str, Any], List[str]]:
    """Deterministic, capped cross-customer queries for the need_data protocol.

    ``branch`` filters by the CUSTOMER's branch (exact, case-insensitive); when
    it matches nothing the full (capped) list is returned WITH a branch column
    plus a warning, so the model can still filter — never silently empty."""
    from sqlalchemy import select

    from app.models.customer import Customer
    from app.models.facility import Facility
    from app.models.guarantor import Guarantor
    from app.models.profile_entities import FixedDeposit, MortgagedProperty
    from app.models.security import Security

    warnings: List[str] = []
    out: Dict[str, Any] = {}
    b = (branch or "").strip()

    custs = (
        await db.execute(select(Customer).where(Customer.is_deleted == False))  # noqa: E712
    ).scalars().all()
    by_acc = {c.account_no: c for c in custs if c.account_no}
    accounts: set | None = None
    if b:
        bl = b.lower()
        matched = {c.account_no for c in custs if _s(c.branch).lower() == bl}
        if matched:
            accounts = matched
        else:
            warnings.append(
                f"مشتری‌ای با شعبهٔ «{b}» یافت نشد — فهرستِ کامل با ستونِ شعبه ارسال شد"
            )

    def keep(acc: str) -> bool:
        return accounts is None or acc in accounts

    def cust_info(acc: str) -> Dict[str, str]:
        c = by_acc.get(acc)
        return {
            "customer_name": _s(getattr(c, "name", "")) if c else "",
            "branch": _s(getattr(c, "branch", "")) if c else "",
            "account_manager": _s(getattr(c, "relationship_manager", "")) if c else "",
        }

    if "properties" in datasets:
        rows = (
            await db.execute(select(MortgagedProperty).where(MortgagedProperty.is_deleted == False))  # noqa: E712
        ).scalars().all()
        items = []
        for p in rows:
            if not keep(p.account_no):
                continue
            ci = cust_info(p.account_no)
            items.append({
                "account_no": _s(p.account_no),
                "customer_name": _s(p.customer_name) or ci["customer_name"],
                "branch": ci["branch"], "account_manager": ci["account_manager"],
                "plate_no": _s(p.plate_no), "mortgage_deed_no": _s(p.mortgage_deed_no),
                "city": _s(p.city), "prop_type": _s(p.prop_type), "owner": _s(p.owner),
                "valuation": _s(p.valuation), "valuation_currency": _s(p.valuation_currency),
                "insurance_no": _s(p.insurance_no), "insurance_expiry": _s(p.insurance_expiry),
                "mortgage_date": _s(p.mortgage_date), "mortgage_amount": _s(p.mortgage_amount),
                "remarks": _s(p.remarks),
                # v110 — the stored-but-unexposed fields + the policy identity
                # block, so the collateral/insurance table can be built from DB
                "owner_national_id": _s(p.owner_national_id), "postal_code": _s(p.postal_code),
                "address": _s(p.address), "land_area": _s(p.land_area),
                "infra_area": _s(p.infra_area), "building_age": _s(p.building_age),
                "zone": _s(p.zone), "last_valuation_date": _s(p.last_valuation_date),
                "insurance_issue": _s(p.insurance_issue),
                "insurance_computer_code": _s(p.insurance_computer_code),
                "insurance_policyholder": _s(p.insurance_policyholder),
                "insurance_subject": _s(p.insurance_subject),
                "insurance_activity": _s(p.insurance_activity),
                "insurance_coverage_total": _s(p.insurance_coverage_total),
                "insurance_issuing_unit": _s(p.insurance_issuing_unit),
            })
        out["properties"] = _cap("املاک رهنی", items, warnings)

    if "customers" in datasets:
        items = [
            {"account_no": _s(c.account_no), "name": _s(c.name),
             "branch": _s(c.branch), "account_manager": _s(c.relationship_manager)}
            for c in custs if keep(c.account_no)
        ]
        out["customers"] = _cap("مشتریان", items, warnings)

    if "facilities" in datasets:
        facs = (
            await db.execute(
                select(Facility, Customer.account_no)
                .join(Customer, Facility.customer_id == Customer.id)
                .where(Facility.is_deleted == False, Customer.is_deleted == False)  # noqa: E712
            )
        ).all()
        items = []
        for fac, acc in facs:
            if not keep(acc):
                continue
            ci = cust_info(acc)
            items.append({
                "account_no": _s(acc), "customer_name": ci["customer_name"],
                "branch": _s(fac.branch) or ci["branch"],
                "account_manager": _s(fac.relationship_manager) or ci["account_manager"],
                "facility": _s(fac.name), "amount": _s(fac.amount), "currency": _s(fac.currency),
                "outstanding": _s(fac.outstanding), "interest_rate": _s(fac.interest_rate),
                "start_date": _s(fac.start_date), "end_date": _s(fac.end_date),
                "expiry_date": _s(fac.expiry_date),
            })
        out["facilities"] = _cap("تسهیلات", items, warnings)

    if "securities" in datasets:
        rows = (
            await db.execute(select(Security).where(Security.is_deleted == False))  # noqa: E712
        ).scalars().all()
        items = []
        for s in rows:
            if not keep(s.account_no):
                continue
            ci = cust_info(s.account_no)
            items.append({
                "year": _s(s.year), "branch": _s(s.branch) or ci["branch"],
                "account_no": _s(s.account_no),
                "customer_name": _s(s.customer_name) or ci["customer_name"],
                "fd": _s(s.fd), "guarantor": _s(s.guarantor), "cheque_no": _s(s.cheque_no),
                "cheque_amount": _s(s.cheque_amount), "property_no": _s(s.property_no),
                "mortgage_aed": _s(s.mortgage_aed), "remarks": _s(s.remarks),
            })
        out["securities"] = _cap("تضامین", items, warnings)

    if "fixed_deposits" in datasets:
        rows = (
            await db.execute(select(FixedDeposit).where(FixedDeposit.is_deleted == False))  # noqa: E712
        ).scalars().all()
        items = []
        for fd in rows:
            if not keep(fd.account_no):
                continue
            ci = cust_info(fd.account_no)
            items.append({
                "account_no": _s(fd.account_no),
                "customer_name": _s(fd.customer_name) or ci["customer_name"],
                "branch": ci["branch"], "fd_number": _s(fd.fd_number),
                "amount": _s(fd.amount), "currency": _s(fd.currency),
                "open_date": _s(fd.open_date), "maturity_date": _s(fd.maturity_date),
                "rate": _s(fd.rate),
            })
        out["fixed_deposits"] = _cap("سپرده‌ها", items, warnings)

    if "guarantors" in datasets:
        rows = (
            await db.execute(select(Guarantor).where(Guarantor.is_deleted == False))  # noqa: E712
        ).scalars().all()
        items = []
        for g in rows:
            if not keep(g.account_no):
                continue
            ci = cust_info(g.account_no)
            items.append({
                "account_no": _s(g.account_no),
                "customer_name": _s(g.customer_name) or ci["customer_name"],
                "branch": _s(g.branch) or ci["branch"],
                "guarantor_name": _s(g.guarantor_name), "guarantor_account": _s(g.guarantor_account),
                "cheque_no": _s(g.cheque_no), "cheque_amount": _s(g.cheque_amount),
                "issuing_bank": _s(g.issuing_bank),
            })
        out["guarantors"] = _cap("ضامن‌ها", items, warnings)

    # Activity LOGS — audit trail + journal/daily-log lines. The SEARCH runs
    # over the WHOLE tables (log_search — no newest-N pre-limit) with the
    # model's optional logs_filter; only the returned rows are capped, and the
    # true totals surface in warnings so a cut is never silent. The branch
    # filter still narrows post-search through each row's account.
    want_audit = "audit_logs" in datasets
    want_journal = "journal_entries" in datasets
    if want_audit or want_journal:
        from app.services.log_search import search_logs

        lf = dict(logs_filter or {})
        lf["scope"] = "both" if (want_audit and want_journal) else ("audit" if want_audit else "journal")
        found = await search_logs(db, lf)
        warnings.extend(found.get("warnings") or [])
        if want_audit:
            items = []
            for r in found.get("audit") or []:
                acc = r.get("account_no") or ""
                if not keep(acc):
                    continue
                r["customer_name"] = cust_info(acc)["customer_name"] if acc else ""
                items.append(r)
            out["audit_logs"] = _cap("لاگِ کلی", items, warnings)
        if want_journal:
            items = []
            for r in found.get("journal") or []:
                acc = r.get("account_no") or ""
                if not keep(acc):
                    continue
                ci = cust_info(acc) if acc else {"customer_name": "", "branch": "", "account_manager": ""}
                r["customer_name"] = r.get("customer_name") or ci["customer_name"]
                r["branch"] = r.get("branch") or ci["branch"]
                items.append(r)
            out["journal_entries"] = _cap("لاگِ کارها", items, warnings)

    return out, warnings


def _clean_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]+', "-", (name or "").strip())
    name = re.sub(r"\s+", " ", name).strip(" .-") or "پیوست"
    return name[:80]


def finalize_filename(filename: str, account_no: str) -> str:
    """Owner rule: a generated file's name = its CONTENT + the account number.
    The model supplies the content part (rule 7); the account number is appended
    HERE, deterministically — never left to the model. Idempotent (skips when
    the account is already in the name), keeps the extension, and does nothing
    for the general/no-account case."""
    acct = (account_no or "").strip()
    if not acct or acct.lower() == "general":
        return filename
    stem, dot, ext = (filename or "").rpartition(".")
    if not dot:
        stem, ext = filename or "", ""
    if acct in stem:
        return filename
    stem = stem.strip()[:70].rstrip(" -")
    return f"{stem} - حساب {acct}" + (f".{ext}" if ext else "")


def parse_spec(raw_text: str) -> Tuple[Dict[str, Any], List[str]]:
    """Parse + CLAMP the model's spec. Raises ValueError when unusable."""
    m = _JSON_RE.search(raw_text or "")
    if not m:
        raise ValueError("no_json")
    try:
        data = json.loads(m.group(0))
    except Exception as exc:  # noqa: BLE001
        raise ValueError("bad_json") from exc
    if not isinstance(data, dict):
        raise ValueError("bad_json")

    kind = str(data.get("kind") or "").strip().lower()
    if kind not in ("excel", "word"):
        # infer: sheets ⇒ excel, paragraphs ⇒ word
        kind = "excel" if data.get("sheets") else "word"

    warnings = [str(w)[:300] for w in (data.get("warnings") or []) if str(w).strip()][:10]
    spec: Dict[str, Any] = {
        "kind": kind,
        "filename": _clean_filename(str(data.get("filename") or "پیوست")),
        "title": str(data.get("title") or "").strip()[:200],
    }

    if kind == "excel":
        sheets_in = data.get("sheets") or []
        if not isinstance(sheets_in, list) or not sheets_in:
            raise ValueError("no_sheets")
        sheets: List[Dict[str, Any]] = []
        for sh in sheets_in[:MAX_SHEETS]:
            if not isinstance(sh, dict):
                continue
            cols = [str(c)[:120] for c in (sh.get("columns") or [])][:MAX_COLS]
            if not cols:
                continue
            rows: List[List[str]] = []
            for row in (sh.get("rows") or [])[:MAX_ROWS]:
                if not isinstance(row, list):
                    continue
                rows.append([("" if c is None else str(c))[:MAX_CELL] for c in row[: len(cols)]])
            name = re.sub(r"[\\/*?\[\]:]+", "-", str(sh.get("name") or "برگه"))[:30] or "برگه"
            sheets.append({"name": name, "columns": cols, "rows": rows})
        if not sheets:
            raise ValueError("no_sheets")
        spec["sheets"] = sheets
    else:
        paras_in = data.get("paragraphs") or []
        if not isinstance(paras_in, list) or not paras_in:
            raise ValueError("no_paragraphs")
        paras: List[Dict[str, Any]] = []
        for p in paras_in[:MAX_PARAGRAPHS]:
            if isinstance(p, str):
                p = {"text": p}
            if not isinstance(p, dict):
                continue
            text = str(p.get("text") or "").strip()
            if not text:
                continue
            align = str(p.get("align") or "").lower()
            paras.append({
                "text": text[:4000],
                "heading": bool(p.get("heading")),
                "bold": bool(p.get("bold")),
                "align": align if align in ("right", "center", "justify") else "justify",
            })
        if not paras:
            raise ValueError("no_paragraphs")
        spec["paragraphs"] = paras

    return spec, warnings


# ---------------------------------------------------------------------------
# Renderers — clean, official-looking output (RTL, Persian-friendly fonts).
# ---------------------------------------------------------------------------

def render_excel(spec: Dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)
    thin = Side(style="thin", color="444444")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="DCE6F1")
    head_font = Font(name="B Nazanin", bold=True, size=12)
    cell_font = Font(name="B Nazanin", size=11)
    title_font = Font(name="B Titr", bold=True, size=13)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center", wrap_text=True)

    for sh in spec["sheets"]:
        ws = wb.create_sheet(title=sh["name"])
        ws.sheet_view.rightToLeft = True
        ncols = len(sh["columns"])
        r0 = 1
        if spec.get("title"):
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(1, ncols))
            tc = ws.cell(row=1, column=1, value=spec["title"])
            tc.font = title_font
            tc.alignment = center
            ws.row_dimensions[1].height = 26
            r0 = 2
        for ci, col in enumerate(sh["columns"], start=1):
            c = ws.cell(row=r0, column=ci, value=col)
            c.font = head_font
            c.fill = head_fill
            c.border = border
            c.alignment = center
        ws.row_dimensions[r0].height = 22
        for ri, row in enumerate(sh["rows"], start=r0 + 1):
            for ci in range(1, ncols + 1):
                val = row[ci - 1] if ci - 1 < len(row) else ""
                c = ws.cell(row=ri, column=ci, value=val)
                c.font = cell_font
                c.border = border
                c.alignment = right
        # column widths sized to content (clamped so the sheet stays printable)
        for ci in range(1, ncols + 1):
            longest = max(
                [len(str(sh["columns"][ci - 1]))]
                + [len(str(r[ci - 1])) if ci - 1 < len(r) else 0 for r in sh["rows"]]
                or [8]
            )
            ws.column_dimensions[get_column_letter(ci)].width = min(60, max(10, longest + 4))
        ws.freeze_panes = ws.cell(row=r0 + 1, column=1)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def render_word(spec: Dict[str, Any]) -> bytes:
    import docx  # python-docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def make_rtl(paragraph) -> None:
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)

    def style_run(run, *, bold: bool, size: int) -> None:
        run.font.name = "B Nazanin"
        run.font.size = Pt(size)
        run.bold = bold
        rPr = run._r.get_or_add_rPr()
        rfonts = rPr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rPr.append(rfonts)
        rfonts.set(qn("w:cs"), "B Nazanin")
        szcs = OxmlElement("w:szCs")
        szcs.set(qn("w:val"), str(size * 2))
        rPr.append(szcs)
        if bold:
            bcs = OxmlElement("w:bCs")
            rPr.append(bcs)

    doc = docx.Document()
    if spec.get("title"):
        p = doc.add_paragraph()
        make_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(spec["title"]), bold=True, size=15)
    for para in spec["paragraphs"]:
        p = doc.add_paragraph()
        make_rtl(p)
        p.alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[para["align"]]
        style_run(
            p.add_run(para["text"]),
            bold=bool(para["bold"] or para["heading"]),
            size=14 if para["heading"] else 12,
        )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(spec: Dict[str, Any]) -> Tuple[bytes, str, str]:
    """Render the validated spec → (bytes, full filename, mimetype)."""
    if spec["kind"] == "excel":
        data = render_excel(spec)
        return data, f"{spec['filename']}.xlsx", (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    data = render_word(spec)
    return data, f"{spec['filename']}.docx", (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
