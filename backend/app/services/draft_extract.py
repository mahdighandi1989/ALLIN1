"""Extract data from a filled credit-committee approval draft (مصوبه / "draft
sanction") Word file, so the Offer Letter form can be prefilled and the
customer's profile enriched.

The draft is mostly tables with consistent labels ("Customer Name:", "Account
Number:", "Branch Name", "Borrower Type", "Purpose:", "Date of Review:", ...) and
prose that states the proposed facility ("fresh personal loan of AED 80,000/- for
a period of 48 months"). We flatten the document to lines, read the "Label: value"
pairs, and pull the rest with a few targeted regexes. Nothing here writes to the
DB — the caller decides what to persist (idempotently).
"""
from __future__ import annotations

import re
from io import BytesIO
from typing import Optional

from docx import Document


def _lines(doc: Document) -> list[str]:
    """Every non-empty text line in reading order (paragraphs + table cells).
    Merged table cells repeat their text; that is harmless for label lookups."""
    out: list[str] = []
    for p in doc.paragraphs:
        for ln in p.text.split("\n"):
            ln = ln.strip()
            if ln:
                out.append(ln)
    for t in doc.tables:
        for row in t.rows:
            prev = None
            for c in row.cells:
                txt = c.text.strip()
                if not txt or txt == prev:
                    continue
                prev = txt
                for ln in txt.split("\n"):
                    ln = ln.strip()
                    if ln:
                        out.append(ln)
    return out


def _kv(lines: list[str]) -> dict[str, str]:
    """Inline "Label: value" pairs (first occurrence wins)."""
    d: dict[str, str] = {}
    pat = re.compile(r"^([A-Za-z][A-Za-z0-9 ()/&#.\-'’]{2,60}?)\s*[:：]\s*(.+)$")
    for ln in lines:
        m = pat.match(ln)
        if m:
            k = re.sub(r"\s+", " ", m.group(1).strip().lower())
            v = m.group(2).strip()
            if v and v not in (":",) and k not in d:
                d[k] = v
    return d


def _after(lines: list[str], *labels: str) -> str:
    """First non-empty, non-label line following a standalone block label
    (e.g. the paragraph under "Purpose:")."""
    wanted = {l.lower().rstrip(":").strip() for l in labels}
    for i, ln in enumerate(lines):
        if ln.lower().rstrip(":").strip() in wanted:
            for j in range(i + 1, min(i + 4, len(lines))):
                nxt = lines[j].strip()
                if nxt and not nxt.endswith(":") and len(nxt) > 3:
                    return nxt
    return ""


def _cell_pairs(doc) -> dict[str, str]:
    """label → value from ADJACENT table cells (the common "Label | value"
    layout), so we catch fields that aren't written as inline "Label: value"."""
    d: dict[str, str] = {}
    for t in doc.tables:
        for row in t.rows:
            seq: list[str] = []
            prev = None
            for c in row.cells:
                tx = c.text.strip()
                if tx and tx != prev:
                    seq.append(tx)
                    prev = tx
            for i in range(len(seq) - 1):
                key = re.sub(r"\s+", " ", seq[i].strip().rstrip(":").strip().lower())
                val = seq[i + 1].strip()
                if 2 <= len(key) <= 60 and val and not key.replace(" ", "").isdigit() and "\n" not in seq[i]:
                    d.setdefault(key, val)
    return d



_BRANCH_NAMES = {
    "2533": "BUR DUBAI", "2690": "ABU DHABI", "2776": "SHARJAH", "2900": "AJMAN",
    "4350": "SHEIKH ZAYED ROAD", "2624": "AL MAKTOUM", "2898": "MURSHID BAZAR",
    "1741": "AL AIN", "3535": "HEAD OFFICE",
}


def _clean(v: str) -> str:
    return re.sub(r"\s+", " ", (v or "").strip()).strip(" .")


def extract_from_docx(data: bytes) -> dict:
    """Parse a draft-sanction .docx into offer-letter + profile fields."""
    doc = Document(BytesIO(data))
    lines = _lines(doc)
    full = "\n".join(lines)
    kv = _kv(lines)
    cellkv = _cell_pairs(doc)

    def g(*labels: str) -> str:
        """Robust label lookup: inline "Label: value" → adjacent cell → block."""
        for lb in labels:
            k = re.sub(r"\s+", " ", lb.strip().lower())
            if kv.get(k):
                return kv[k].strip()
        for lb in labels:
            k = re.sub(r"\s+", " ", lb.strip().lower())
            if cellkv.get(k):
                return cellkv[k].strip()
        return _after(lines, *labels)

    # --- identity ---------------------------------------------------------
    name = _clean(g("customer name", "customer's name", "customer’s name", "name of customer", "borrower name", "name of borrower"))
    acc_raw = g("account number", "account no", "account no.", "a/c no", "a/c no.", "account #", "acct no", "cif")
    nums = re.findall(r"\d+", acc_raw)
    account_no = next((n for n in nums if len(n) == 6), (nums[1] if len(nums) > 1 else (nums[0] if nums else "")))
    branch_code = nums[0] if nums and len(nums[0]) == 4 else ""
    suffix = nums[-1] if len(nums) >= 3 else ""
    account_display = "-".join(nums) if nums else acc_raw

    branch_name = _clean(g("branch name", "branch", "branch name & code", "branch name and code", "branch & code", "branch code"))
    bc = branch_code or (re.search(r"\b(\d{4})\b", branch_name).group(1) if re.search(r"\b(\d{4})\b", branch_name) else "")
    branch_label = f"{_BRANCH_NAMES.get(bc)} - {bc}" if bc and _BRANCH_NAMES.get(bc) else branch_name

    btype = (g("borrower type", "customer type", "type of borrower", "segment") or "").lower()
    account_type = "corporate" if ("corporate" in btype or "sme" in btype) else ("retail" if ("retail" in btype or "individual" in btype) else "")

    # --- proposed facility: amount / tenor / rate / type (independent) -----
    amount = ""
    for pat in (
        r"fresh\s+(?:commercial|personal)\s+loan[^.\n]*?AED\s*([\d,]+)",
        r"(?:proposed|fresh)[^.\n]*?loan[^.\n]*?AED\s*([\d,]+)",
        r"loan (?:facility )?of\s*AED\s*([\d,]+)",
        r"AED\s*([\d,]+)\s*/?-?\s*for a period of",
    ):
        mm = re.search(pat, full, re.I)
        if mm:
            amount = mm.group(1).replace(",", "")
            break
    mt = (re.search(r"period of\s*(\d+)\s*months", full, re.I)
          or re.search(r"(\d+)\s*months?\s*from\s*dod", full, re.I)
          or re.search(r"proposed[^\n]*?(\d{1,2})\s*months", full, re.I))
    tenor = mt.group(1) if mt else ""
    mf = re.search(r"fresh\s+(commercial|personal)\s+loan", full, re.I)
    if mf:
        facility = f"{mf.group(1).title()} Loan"
    elif re.search(r"\bPIM\b|personal loan", full, re.I):
        facility = "Personal Loan"
    elif re.search(r"commercial loan|\bCL[\b-]", full, re.I):
        facility = "Commercial Loan"
    else:
        facility = ""
    mr = (re.search(r"to be(?:\s+at)?\s*([\d.]+)\s*%\s*p\.?\s*a", full, re.I)
          or re.search(r"([\d.]+)\s*%\s*(?:p\.?\s*a|per annum)", full, re.I))
    rate = f"{mr.group(1)}% p.a." if mr else ""

    purpose = _clean(_after(lines, "Purpose") or g("purpose"))
    # The customer's application/request-letter date → the Offer Letter "Subject" date.
    mapp = re.search(r"customer (?:request )?letter (?:dated|dtd)\.?\s*(\d{1,2}[./-]\d{1,2}[./-]\d{2,4})", full, re.I)
    app_date = mapp.group(1) if mapp else ""
    business = _clean(g("business activity", "business type", "nature of business", "activity"))
    rating = _clean(g("proposed customer rating", "proposed rating", "customer rating", "rating"))
    existing_rating = _clean(g("existing customer rating", "existing rating"))

    # --- profile enrichment ----------------------------------------------
    tl_no = tl_exp = ""
    mtl = re.search(r"trade licen[cs]e no\.?\s*([A-Za-z0-9\-]+).*?(?:till|validity|valid till|expiry)[^\d]*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", full, re.I)
    if mtl:
        tl_no, tl_exp = mtl.group(1), mtl.group(2)
    if not tl_no:
        tl_no = _clean(g("trade license no", "trade license", "trade licence no", "license no"))

    established = _clean(g("company/ firm established since", "company/firm established since", "established since", "date of establishment", "incorporation date", "establishment date"))
    if not established:
        me = re.search(r"established on\s*([0-9A-Za-z/\-]+)", full, re.I)
        established = me.group(1) if me else ""
    relationship = _clean(g("relationship date", "relationship since", "banking relationship since"))
    auditor = _clean(g("auditor’s name", "auditor's name", "auditor", "audited by"))
    credit_app = _clean(g("credit application #", "loan application #", "credit application no", "loan application no", "application #", "ca #"))
    review_date = _clean(g("date of review", "review date"))

    ma = re.search(r"credit score is[^\d]*([\d]{3,4})", full, re.I) or re.search(r"\baecb\b[^\d]{0,20}(\d{3,4})", full, re.I)
    aecb = ma.group(1) if ma else ""
    mad = re.search(r"located at\s*(.+?)(?:\.\s|\.$|\n)", full, re.I)
    address = _clean(mad.group(1)) if mad else _clean(g("address", "office address", "registered address"))
    ms = re.search(r"monthly salary of AED\s*([\d,]+)", full, re.I) or re.search(r"salary[^\d]{0,20}AED\s*([\d,]+)", full, re.I)
    salary = ms.group(1).replace(",", "") if ms else ""
    profile_text = _clean(_after(lines, "Customer Profile"))

    # --- guarantors (rows like "Name (Guarantor -1) | 2624 | 124076") -----
    guarantors = []
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            joined = " | ".join(cells)
            if re.search(r"\(\s*guarantor", joined, re.I):
                gname = re.sub(r"\(.*?\)", "", cells[0]).strip()
                gnums = re.findall(r"\b\d{4,6}\b", joined)
                gbranch = next((n for n in gnums if len(n) == 4), "")
                gacc = next((n for n in gnums if len(n) == 6), "")
                if gname:
                    guarantors.append({"name": _clean(gname), "branch": gbranch, "account": gacc})

    offer = {
        "CompanyName": name,
        "AccountNumber": account_display or account_no,
        "AccountType": account_type,
        "Branch": branch_label,
        "FacilityType": facility,
        "Purpose": purpose,
        "BusinessType": business,
        "Rating": rating,
        "ExistingRating": existing_rating,
        "LoanTenor": tenor,
        "InterestRate": rate,
        "LoanInterestRate": rate,
        "LoanAmount": f"{int(amount):,}" if amount.isdigit() else "",
        "CreditLimit": f"{int(amount):,}" if amount.isdigit() else "",
        "SubjectDate": app_date or review_date,
    }
    profile = {
        "trade_license_no": tl_no,
        "trade_license_expiry": tl_exp,
        "business_type": business,
        "established_since": established,
        "relationship_date": relationship,
        "auditor": auditor,
        "credit_application_no": credit_app,
        "review_date": review_date,
        "aecb_score": aecb,
        "address": address,
        "monthly_salary": salary,
        "customer_profile": profile_text,
        "proposed_amount": amount,
        "proposed_tenor": tenor,
        "proposed_facility": facility,
        "proposed_rate": rate,
    }
    # Drop empties so we never overwrite good data with blanks.
    offer = {k: v for k, v in offer.items() if v}
    profile = {k: v for k, v in profile.items() if v}
    return {
        "account_no": account_no,
        "account_display": account_display,
        "branch_code": bc,
        "suffix": suffix,
        "offer": offer,
        "profile": profile,
        "guarantors": guarantors,
    }
